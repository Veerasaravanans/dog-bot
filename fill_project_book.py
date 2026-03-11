"""
Script to fill Synopsis/Abstract and Review Summary sections
in the Student's Project Book Report.docx
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import shutil

DOCX_PATH = os.path.join('project report', 'Student\u2019s Project Book Report.docx')
BACKUP_PATH = DOCX_PATH.replace('.docx', '_backup.docx')

# ── Content ──────────────────────────────────────────────────────────────────

SYNOPSIS_ABSTRACT = (
    "The DogBot Recon System is an autonomous reconnaissance robot engineered for intelligent "
    "navigation and real-time obstacle detection in unstructured indoor environments. The growing "
    "demand for autonomous mobile robots in surveillance, search-and-rescue, and hazardous area "
    "inspection necessitates systems that combine affordability with robust perception and decision-making "
    "capabilities. This project addresses that need by integrating an ESP32-CAM microcontroller for visual "
    "perception with a layered software architecture running on a FastAPI-based backend server."
    "\n\n"
    "The vision subsystem employs a multi-sensor fusion approach combining four independent computer vision "
    "algorithms\u2014MOG2 background subtraction, Canny edge detection, Sobel gradient magnitude analysis, and "
    "dense Farneback optical flow\u2014requiring agreement from at least two sensors per pixel to confirm obstacle "
    "presence. This fusion strategy significantly reduces false positives while maintaining high detection "
    "sensitivity. Perspective-based distance estimation maps pixel coordinates to real-world distances using "
    "inverse perspective transformation calibrated to the camera\u2019s 15 cm mounting height. Auxiliary detection "
    "layers include thin-edge detection for narrow vertical structures such as table legs and poles, and "
    "adaptive floor-color modeling using HSV statistics with exponential moving average smoothing."
    "\n\n"
    "Object detection and classification are handled by a YOLOv8n-seg segmentation model with ByteTrack "
    "multi-object tracking, providing instance-level segmentation masks and stable tracking identifiers across "
    "frames. The detection pipeline runs on a configurable interval to balance accuracy against computational "
    "cost, with results classified into DANGER, CAUTION, and SAFE zones based on vertical position in the frame."
    "\n\n"
    "The core navigation engine implements a five-stage real-time path planning pipeline executing in "
    "approximately 2 milliseconds per frame: (1) occupancy grid construction projecting detected obstacles "
    "into a 60\u00d780-cell bird\u2019s-eye grid at 5 cm resolution covering 3.0 m \u00d7 4.0 m, (2) Kalman filter-based "
    "obstacle state estimation predicting position and velocity, (3) multi-layer cost map generation combining "
    "occupancy, repulsive potential fields, predictive hazard projections, boundary penalties, and forward-progress "
    "attraction, (4) trajectory evaluation sampling 45 candidate arcs across 15 steering angles and 3 speed "
    "levels scored by obstacle proximity, heading deviation, smoothness, clearance, and forward progress, and "
    "(5) temporal smoothing via exponential moving average to prevent control jitter."
    "\n\n"
    "The AI decision engine operates on a Stop-Analyze-Act cycle with a 5-frame direction voting mechanism "
    "requiring 60% majority consensus, ensuring robust directional decisions resistant to single-frame noise. "
    "A VIO Cloud LLM serves as an asynchronous fallback for low-confidence scenarios. The system supports "
    "Manual, Semi-Autonomous, and Search control modes with seamless transitions."
    "\n\n"
    "Dual connectivity is achieved through local HTTP communication for same-network operation and MQTT with "
    "TLS encryption via HiveMQ Cloud for remote operation across networks. The ESP32-CAM firmware implements "
    "L298N motor control with PWM speed regulation, automatic drift compensation, and an 800 ms failsafe "
    "auto-stop mechanism. The complete system processes frames at 15 FPS with WebSocket streaming for live "
    "video and bidirectional control, demonstrating viable autonomous navigation in real-world test environments "
    "using cost-effective hardware."
)

SYNOPSIS_FINAL_DRAFT = (
    "The DogBot Recon System is an autonomous reconnaissance robot designed for intelligent navigation "
    "and real-time obstacle detection in complex indoor environments. The system integrates an ESP32-CAM "
    "microcontroller for visual perception with a FastAPI-based backend implementing advanced computer vision "
    "and AI-powered decision-making algorithms."
    "\n\n"
    "The vision subsystem employs multi-sensor fusion combining four independent algorithms\u2014MOG2 background "
    "subtraction, Canny edge detection, Sobel gradient magnitude, and dense Farneback optical flow\u2014with a "
    "minimum two-sensor agreement threshold to confirm obstacles, minimizing false positives while preserving "
    "detection sensitivity. Perspective-based distance estimation using inverse perspective transformation, "
    "calibrated at 15 cm camera height, maps pixel positions to real-world distances. Supplementary detection "
    "includes thin-edge detection for narrow vertical structures and adaptive floor-color modeling using HSV "
    "statistics with EMA smoothing."
    "\n\n"
    "Object detection utilizes YOLOv8n-seg with ByteTrack multi-object tracking for instance segmentation "
    "and persistent object identification. Detected objects are classified into DANGER, CAUTION, and SAFE "
    "zones for prioritized threat assessment."
    "\n\n"
    "Navigation is governed by a five-stage real-time path planning pipeline (~2 ms latency): occupancy grid "
    "construction (60\u00d780 cells, 5 cm resolution, 3.0 m \u00d7 4.0 m coverage), Kalman filter-based obstacle state "
    "estimation, multi-layer cost map generation (occupancy, repulsive, predictive, boundary, and attractive "
    "components), trajectory evaluation across 45 candidate arcs scored on five cost criteria, and temporal "
    "smoothing via exponential moving average. Emergency stop engages at 0.25 m with adaptive recovery behavior "
    "after sustained path blockage."
    "\n\n"
    "The AI decision engine implements a Stop-Analyze-Act cycle with 5-frame direction voting (60% majority "
    "required), supported by VIO Cloud LLM fallback for low-confidence decisions. Control modes include Manual, "
    "Semi-Autonomous, and Search with seamless transitions."
    "\n\n"
    "Dual connectivity through HTTP (local) and TLS-encrypted MQTT via HiveMQ Cloud (remote) enables operation "
    "across network boundaries. The ESP32-CAM firmware provides L298N motor control with PWM speed regulation, "
    "drift compensation, and 800 ms failsafe auto-stop. A web dashboard with real-time WebSocket streaming "
    "provides live video, telemetry, and interactive controls. The system achieves 15 FPS processing throughput "
    "and demonstrates effective autonomous navigation using cost-effective embedded hardware."
)

REVIEW_1_SUMMARY = (
    "The first review established the foundation of the DogBot Recon System through comprehensive project "
    "planning, system architecture design, and technology selection. The project scope was defined to develop "
    "an autonomous reconnaissance robot capable of intelligent navigation, real-time obstacle detection, and "
    "remote operation in unstructured indoor environments."
    "\n\n"
    "A thorough literature survey was conducted covering autonomous mobile robotics, embedded vision systems, "
    "real-time path planning algorithms, and AI-based decision-making frameworks. The survey identified key "
    "gaps in affordable autonomous navigation systems that combine embedded hardware with advanced software "
    "pipelines, motivating the project\u2019s approach."
    "\n\n"
    "The system architecture was designed as a service-oriented framework with clearly defined layers: an "
    "ESP32-CAM microcontroller for visual perception and motor actuation, a FastAPI-based backend server for "
    "real-time frame processing and decision-making, and WebSocket-based communication for live video streaming "
    "and bidirectional control. The hardware platform was selected as the AI-Thinker ESP32-CAM module paired "
    "with an L298N motor driver, chosen for their balance of cost-effectiveness, computational capability, and "
    "community support."
    "\n\n"
    "Key software framework decisions included FastAPI for asynchronous web services, OpenCV for computer vision "
    "processing, Ultralytics YOLOv8 for object detection, and MQTT with TLS encryption for remote connectivity "
    "via HiveMQ Cloud. The development environment was configured with Python virtual environments, and the "
    "ESP32 firmware development workflow was established using the Arduino IDE with ESP32 board support."
    "\n\n"
    "Initial system requirements were documented specifying 15 FPS target processing rate, real-time obstacle "
    "detection with distance estimation, autonomous path planning, multiple control modes (Manual, Semi-Autonomous), "
    "and dual connectivity supporting both local HTTP and remote MQTT operation. The project activity plan and "
    "Gantt chart were prepared, outlining the development schedule across four review milestones."
)

REVIEW_2_SUMMARY = (
    "The second review demonstrated the successful implementation and integration of core hardware and software "
    "components of the DogBot Recon System. The ESP32-CAM firmware was developed and deployed, establishing "
    "dual-server operation with an HTTP control server on port 80 and an MJPEG video streaming server on port 81. "
    "Camera initialization was configured for VGA resolution (640\u00d7480) with optimized JPEG quality and dual frame "
    "buffers utilizing PSRAM."
    "\n\n"
    "The L298N motor driver integration was completed with PWM speed control on ENA/ENB pins, enabling variable "
    "speed operation. Motor control functions for forward, backward, left turn, and right turn were implemented "
    "with a firmware-level motor trim constant to compensate for observed rightward drift during straight-line "
    "travel. The failsafe auto-stop mechanism was programmed to halt motors after 800 milliseconds of command "
    "inactivity, preventing runaway scenarios."
    "\n\n"
    "The computer vision pipeline was implemented as the core perception module, incorporating four independent "
    "detection algorithms: MOG2 background subtraction with 300-frame history and shadow detection, Canny edge "
    "detection with configurable thresholds (low=80, high=200), Sobel gradient magnitude analysis for "
    "texture-agnostic obstacle boundaries, and dense Farneback optical flow for motion-based detection. The "
    "multi-sensor fusion strategy was implemented requiring minimum two-sensor agreement per pixel region, "
    "significantly reducing false positive detections. Perspective-based distance estimation was calibrated "
    "using inverse perspective mapping with the camera mounted at 15 cm height and a vanishing point fraction "
    "of 0.30, enabling distance measurements from 0.1 m to 10 m."
    "\n\n"
    "The path planning engine\u2019s first three stages were completed: the occupancy grid module projecting tracked "
    "obstacles into a 60\u00d780-cell bird\u2019s-eye grid at 5 cm resolution with temporal decay and Gaussian inflation; "
    "the obstacle state estimator implementing Kalman filtering with a constant-velocity model for predicting "
    "obstacle position and velocity; and the cost map generator combining occupancy, repulsive potential fields, "
    "predictive hazard projections, boundary penalties, and forward-progress attraction with configurable weights."
    "\n\n"
    "WebSocket communication was verified for both video streaming (/ws/video) and bidirectional control "
    "(/ws/control) channels. The backend service orchestration was established through the FrameManager, "
    "coordinating the processing pipeline at the target 15 FPS throughput. System testing confirmed stable "
    "real-time frame acquisition, processing, and broadcasting to connected dashboard clients."
)

REVIEW_3_SUMMARY = (
    "The third review marked the completion of advanced perception, decision-making, and full system integration "
    "of the DogBot Recon System. The ML detection module was finalized with YOLOv8n-seg (Nano Segmentation model) "
    "providing instance-level object detection and segmentation masks. ByteTrack multi-object tracking was "
    "integrated for persistent object identification across frames with stable tracking IDs. Detection results "
    "are classified into three threat zones\u2014DANGER (bottom 25%), CAUTION (25\u201350%), and SAFE (top 50%)\u2014based on "
    "vertical frame position, enabling distance-aware threat prioritization."
    "\n\n"
    "The trajectory evaluator, completing the five-stage path planning pipeline, was implemented and integrated. "
    "The evaluator samples 45 candidate trajectories across 15 steering angles (\u22121.0 to +1.0) and 3 speed levels "
    "(0.33, 0.66, 1.0), generating circular arc paths with a minimum turning radius of 0.30 m and a maximum arc "
    "angle of 75 degrees. Each trajectory is scored using five weighted cost criteria: obstacle proximity (0.30), "
    "heading deviation (0.20), control smoothness (0.10), lateral clearance (0.15), and forward progress (0.25). "
    "Feasibility filters reject trajectories exceeding 0.9 cost per waypoint or falling below 8 cm minimum "
    "clearance. Temporal smoothing via exponential moving average (\u03b1=0.2) prevents control jitter between "
    "consecutive frames."
    "\n\n"
    "The AI decision engine was deployed with the Stop-Analyze-Act cycle operating at 15 FPS. The analysis phase "
    "accumulates direction votes over approximately 45 frames (3 seconds) into a rolling deque, with the "
    "determination phase requiring 60% majority consensus for directional commands. A forward-recovery mechanism "
    "biases toward straight travel after consecutive identical turns if the path planner allocates at least 30% "
    "votes to forward motion. The VIO Cloud LLM integration was implemented as an asynchronous background task "
    "invoked only when local planner confidence falls below 0.4, with 10-second result caching to minimize "
    "API latency impact on navigation."
    "\n\n"
    "MQTT remote connectivity was tested end-to-end with TLS-encrypted communication to the HiveMQ Cloud broker "
    "on port 8883. The ESP32 firmware publishes heartbeat and status telemetry and subscribes to motor command "
    "topics with JSON payload parsing. The backend MQTT bridge was validated for reliable message delivery across "
    "network boundaries."
    "\n\n"
    "Auxiliary perception features were finalized: thin-edge detection using Sobel gradients with orientation "
    "filtering (|grad_x| > 2\u00d7|grad_y|) for detecting narrow vertical structures; adaptive floor-color "
    "detection using HSV per-channel statistics with EMA smoothing (\u03b1=0.05) and \u00b12.5\u03c3 deviation bands; and "
    "vanishing point auto-detection via probabilistic Hough line intersections updated every 30 frames. "
    "Multi-frame obstacle tracking with centroid-based assignment, velocity estimation, and 5-level threat "
    "classification (incorporating distance, time-to-collision, and approach velocity) was integrated into "
    "the CV pipeline. Full end-to-end integration testing demonstrated seamless coordination across all "
    "service components."
)

REVIEW_4_SUMMARY = (
    "The fourth review concluded the DogBot Recon System development with system optimization, robustness "
    "hardening, platform compatibility resolution, and comprehensive documentation. Emergency stop protocols "
    "were implemented with a 0.25 m critical safety threshold in the forward corridor (\u00b10.18 m lateral), "
    "triggering immediate motor halt when obstacles breach the safety boundary. Adaptive recovery behavior "
    "was developed to engage after 5 consecutive blocked frames, steering the robot toward the lower-cost "
    "side at 50% speed with 80% steering magnitude until a viable path is recovered."
    "\n\n"
    "A critical Windows platform compatibility issue was identified and resolved involving the asyncio event "
    "loop policy. Windows\u2019 default ProactorEventLoop does not support the socket-level add_reader/add_writer "
    "operations required by aiomqtt for MQTT communication. The solution implemented in run.py explicitly "
    "creates a SelectorEventLoop and runs the uvicorn server within it using loop.run_until_complete(), "
    "ensuring all asynchronous operations including MQTT execute in a compatible event loop. This fix was "
    "validated across Windows, Linux, and macOS environments."
    "\n\n"
    "Failsafe mechanisms were comprehensively validated: automatic webcam fallback when the ESP32-CAM stream "
    "is unreachable, graceful MQTT reconnection with 5-second retry intervals on connection loss, 800 ms "
    "firmware-level motor auto-stop preventing runaway on controller disconnect, and backpressure handling "
    "in WebSocket frame broadcasting to prevent client lag accumulation. The motor drift compensation system "
    "was calibrated with a firmware trim value of 15 and a backend steering bias of +0.05 to correct observed "
    "rightward drift during straight-line navigation."
    "\n\n"
    "The web-based control dashboard was finalized with real-time features including live 640\u00d7480 video with "
    "HUD overlay (FPS, latency, mode, detection count), perspective grid overlay with distance markers at "
    "0.5 m, 1.0 m, and 2.0 m, interactive D-pad controls for manual motor commands, mode switching between "
    "Manual, Semi-Autonomous, and Search modes, and mini-map visualization of detected obstacle trails. "
    "The REST API was completed with endpoints for health monitoring, telemetry retrieval, mode control, "
    "and decision/detection history."
    "\n\n"
    "Performance benchmarking confirmed the system achieves the target 15 FPS processing throughput with "
    "approximately 2 ms path planning latency per frame. The complete processing pipeline\u2014frame acquisition, "
    "CV processing, ML detection, AI decision-making, path planning, annotation, and WebSocket broadcast\u2014"
    "operates within the 67 ms per-cycle budget. Comprehensive project documentation was completed including "
    "architecture overview, hardware wiring reference, troubleshooting guide, remote access setup guide, and "
    "deployment procedures. The system achieved production-ready status with demonstrated autonomous "
    "navigation capabilities in real-world indoor test environments using cost-effective embedded hardware."
)

# ── Insertion Logic ──────────────────────────────────────────────────────────

def set_paragraph_text(para, text):
    """Set text on a paragraph, preserving its style."""
    # Clear existing runs
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)


def fill_section(doc, start_idx, text, justify=True):
    """Write text into the first empty Body Text paragraph at start_idx."""
    para = doc.paragraphs[start_idx]
    set_paragraph_text(para, text)
    if justify:
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def main():
    # Backup
    if not os.path.exists(BACKUP_PATH):
        shutil.copy2(DOCX_PATH, BACKUP_PATH)
        print(f"Backup created: {BACKUP_PATH}")
    else:
        print(f"Backup already exists: {BACKUP_PATH}")

    doc = Document(DOCX_PATH)

    # 1. Synopsis/Abstract (P[110] is first empty Body Text after "Topic Finalization on")
    print("Filling Synopsis/Abstract (P[110])...")
    fill_section(doc, 110, SYNOPSIS_ABSTRACT)

    # 2. Synopsis/Abstract Final Draft (P[379])
    print("Filling Synopsis/Abstract Final Draft (P[379])...")
    fill_section(doc, 379, SYNOPSIS_FINAL_DRAFT)

    # 3. First Review Summary (P[180] is first empty Body Text after "FIRST REVIEW on")
    print("Filling First Review Summary (P[180])...")
    fill_section(doc, 180, REVIEW_1_SUMMARY)

    # 4. Second Review Summary (P[231])
    print("Filling Second Review Summary (P[231])...")
    fill_section(doc, 231, REVIEW_2_SUMMARY)

    # 5. Third Review Summary (P[279])
    print("Filling Third Review Summary (P[279])...")
    fill_section(doc, 279, REVIEW_3_SUMMARY)

    # 6. Fourth Review Summary (P[329])
    print("Filling Fourth Review Summary (P[329])...")
    fill_section(doc, 329, REVIEW_4_SUMMARY)

    # Save
    doc.save(DOCX_PATH)
    print(f"\nAll sections filled successfully!")
    print(f"File saved: {DOCX_PATH}")
    print(f"\nReminder: Fill 'Topic Finalization on' date manually in the Synopsis/Abstract section.")


if __name__ == '__main__':
    main()
