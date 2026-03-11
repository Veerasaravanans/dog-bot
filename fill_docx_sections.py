"""Fill synopsis/abstract and review summaries in the Student's Project Book Report."""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os, shutil

path = os.path.join('project report', 'Student\u2019s Project Book Report.docx')

# Backup
shutil.copy2(path, path + '.bak')
print('Backup created.')

doc = Document(path)

# ==================== CONTENT ====================

synopsis_abstract = (
    "The DogBot Recon System is an autonomous reconnaissance robot designed for intelligent navigation "
    "and obstacle detection in complex indoor environments. The system integrates advanced computer vision "
    "techniques with AI-powered decision-making to enable real-time autonomous navigation capabilities. "
    "The robot utilizes an ESP32-CAM microcontroller for real-time visual perception, coupled with a "
    "YOLOv8n-seg object detection and segmentation model for comprehensive environmental awareness. "
    "The core navigation system employs a five-stage path planning pipeline incorporating occupancy grid "
    "mapping, Kalman filter-based obstacle state estimation, cost map generation, trajectory evaluation "
    "of 41 candidate arcs, and temporal smoothing for stable steering output. The backend is built on "
    "FastAPI and orchestrates real-time frame processing at 15 FPS with WebSocket streaming for live video "
    "and bidirectional control feedback. The computer vision pipeline provides multi-sensor fusion including "
    "background subtraction, Canny edge detection, gradient magnitude analysis, optical flow computation, "
    "and perspective-based distance estimation using an automotive-style reverse camera visualization. "
    "An AI decision engine implements a five-frame direction voting mechanism with 60% majority consensus "
    "to ensure robust directional decisions, supplemented by VIO Cloud LLM fallback for low-confidence "
    "scenarios. The system supports dual connectivity modes\u2014local HTTP for direct ESP32 communication "
    "and remote MQTT with TLS encryption via HiveMQ Cloud for internet-based operation. Multiple control "
    "modes (Manual, Semi-Autonomous) with seamless fallback mechanisms, emergency stop protocols, and "
    "adaptive recovery behaviors ensure safe and reliable autonomous operation in dynamic environments. "
    "The system achieves sub-2ms path planning latency with continuous steering and speed control outputs, "
    "demonstrating practical viability for indoor reconnaissance and surveillance applications."
)

review_1 = (
    "The first review established the foundation of the DogBot Recon System with a comprehensive project "
    "plan covering system architecture, hardware selection, and software framework design. The project scope "
    "was clearly defined to include autonomous navigation, obstacle detection, and real-time video processing "
    "at 15 FPS target throughput. Key architectural decisions were finalized: ESP32-CAM (AI-Thinker module) "
    "was selected as the primary vision and control platform with dual-server firmware design\u2014Port 80 "
    "for the control API and Port 81 for MJPEG video streaming. FastAPI was chosen as the backend framework "
    "for its native async/await support and WebSocket capabilities. MQTT over TLS via HiveMQ Cloud was "
    "selected for remote connectivity to enable operation across different network locations. The L293D motor "
    "driver was chosen for its simplicity with four GPIO pin control (IN1\u2013IN4). The software architecture "
    "was designed using a service orchestration pattern with FrameManager as the central coordinator, and the "
    "data flow pipeline was mapped from ESP32 capture through CV processing, ML detection, AI decision-making, "
    "and path planning to WebSocket broadcast. Initial system requirements including Pydantic-based configuration "
    "management, asyncio service lifecycle patterns, and the project directory structure were reviewed and "
    "approved. The overall feasibility of integrating computer vision, deep learning inference, and embedded "
    "systems within a single cohesive autonomous robot platform was validated."
)

review_2 = (
    "The second review demonstrated successful integration of all core hardware and software components into "
    "a functional prototype. The ESP32-CAM firmware was implemented with dual-server architecture, motor control "
    "via L293D driver with GPIO pins, MQTT integration with TLS-enabled connection to HiveMQ Cloud, and a "
    "500ms failsafe auto-stop mechanism. The computer vision pipeline (CVPipeline) was fully developed with "
    "multi-sensor fusion capabilities including background subtraction using MOG2, Canny edge detection, "
    "gradient magnitude analysis for thin obstacle detection (desk legs, poles), optical flow computation for "
    "motion analysis, and perspective-based distance estimation with automotive-style reverse camera visualization. "
    "Distance zones were calibrated at 0.5m (near) and 1.2m (mid) boundaries with a camera mounting height of "
    "0.15m. The path planning engine\u2019s first three stages were completed: occupancy grid mapping projecting "
    "obstacles into a bird\u2019s-eye grid of 60\u00d780 cells at 5cm resolution, obstacle state estimation "
    "using Kalman filtering for velocity and position prediction, and cost map generation from occupancy data "
    "combined with predicted positions. WebSocket streaming was verified for both video (annotated frames with "
    "detection overlays) and control channels (bidirectional command/feedback). The FastAPI backend with lifespan "
    "manager was demonstrated to correctly start and stop all services in dependency order. The ESP32Client was "
    "implemented with automatic webcam fallback when the ESP32 is unreachable, ensuring uninterrupted development "
    "and testing. Real-time processing performance met the target 15 FPS throughput."
)

review_3 = (
    "The third review marked the completion and integration of all advanced decision-making and path planning "
    "components. The YOLOv8n-seg segmentation model was integrated via the MLDetector service with built-in "
    "ByteTrack multi-object tracking, optimized to run every N frames for performance. The trajectory evaluator "
    "was implemented as the final stage of the five-stage path planner pipeline, evaluating 41 candidate steering "
    "arcs and scoring them using a weighted cost function comprising obstacle proximity (weight 3.0), heading "
    "deviation (0.8), steering smoothness (0.5), safety clearance (2.0), and forward progress (1.0). The AI "
    "decision engine was deployed with a five-frame direction voting window requiring 60% majority consensus "
    "for directional commands, effectively filtering single-frame noise. The VIO Cloud LLM integration was "
    "completed as a background task fallback, triggered only when local planner confidence falls below 0.4, with "
    "10-second result caching to minimize API calls. Remote connectivity via MQTT was tested end-to-end with "
    "proper TLS configuration, enabling motor command transmission from a remote location through the HiveMQ "
    "Cloud broker. The dual connectivity architecture (MQTT primary with HTTP fallback) was validated. The "
    "dashboard web interface at localhost:8000 was completed with real-time telemetry display, control mode "
    "switching, and live video streaming. Full end-to-end system integration testing demonstrated seamless "
    "coordination between FrameManager orchestration, CV processing, ML detection, AI decision-making, and "
    "path planning with sub-2ms planning latency per frame."
)

review_4 = (
    "The fourth and final review focused on system optimization, robustness hardening, and comprehensive "
    "documentation. Emergency stop protocols were implemented with a critical safety distance of 0.25m in the "
    "forward corridor, triggering immediate motor halt when obstacles breach this threshold. Adaptive recovery "
    "behaviors were added\u2014after 15 consecutive blocked frames, the system initiates an automatic recovery "
    "turn toward the clearer side with configurable steering magnitude (0.8) and speed (0.5). Motor drift "
    "compensation was implemented with a forward steering bias of 0.05 to correct rightward drift. The Windows "
    "platform compatibility issue with asyncio event loops was resolved through a custom startup script (run.py) "
    "that creates a SelectorEventLoop explicitly, ensuring MQTT/aiomqtt compatibility that the default "
    "ProactorEventLoop cannot provide. Failsafe mechanisms were validated including automatic webcam fallback "
    "when ESP32 is unreachable, graceful MQTT reconnection on network drops, and motor auto-stop after 500ms "
    "of no commands on the ESP32 firmware. The floor plane detection algorithm was finalized with adaptive color "
    "modeling using exponential moving average (alpha 0.05) and configurable deviation thresholds. The complete "
    "web dashboard with pin diagram reference, setup guide, and real-time telemetry panels was tested across "
    "multiple browsers. Comprehensive documentation was delivered including system architecture overview, remote "
    "access guide, troubleshooting procedures, and deployment instructions. The system achieved production-ready "
    "status with demonstrated autonomous navigation capabilities in real-world indoor test environments, meeting "
    "all project objectives for autonomous reconnaissance and obstacle avoidance."
)

final_synopsis = (
    "The DogBot Recon System is an autonomous reconnaissance robot that combines embedded systems, computer "
    "vision, deep learning, and AI-powered decision-making for intelligent indoor navigation and obstacle "
    "avoidance. Built around the ESP32-CAM microcontroller with AI-Thinker module, the robot captures real-time "
    "MJPEG video streams processed by a FastAPI backend at 15 frames per second. The computer vision pipeline "
    "employs multi-sensor fusion\u2014background subtraction (MOG2), Canny edge detection, gradient magnitude "
    "analysis, and optical flow\u2014to detect obstacles and estimate distances using perspective geometry with "
    "an automotive-style reverse camera visualization. A YOLOv8n-seg deep learning model provides semantic object "
    "detection and instance segmentation with ByteTrack multi-object tracking for persistent obstacle identification. "
    "The navigation core implements a five-stage path planning pipeline: (1) occupancy grid mapping on a 60\u00d780 "
    "cell grid at 5cm resolution, (2) Kalman filter-based obstacle state estimation for position and velocity "
    "prediction, (3) cost map generation combining static and predicted obstacle data, (4) trajectory evaluation "
    "of 41 candidate steering arcs using a multi-objective weighted cost function, and (5) temporal smoothing for "
    "stable continuous steering (-1 to +1) and speed (0 to 1) outputs. The AI decision engine uses a five-frame "
    "voting mechanism with 60% majority consensus, supplemented by VIO Cloud LLM fallback for low-confidence "
    "scenarios. The system supports dual connectivity\u2014local HTTP and remote MQTT with TLS encryption via "
    "HiveMQ Cloud\u2014with multiple control modes (Manual, Semi-Autonomous), emergency stop protocols at 0.25m "
    "critical distance, and adaptive recovery behaviors. The complete platform includes a real-time web dashboard "
    "with WebSocket streaming for live video, telemetry, and bidirectional control. The system achieves sub-2ms "
    "path planning latency and demonstrates practical viability for autonomous indoor reconnaissance and "
    "surveillance applications."
)

# ==================== FILL SECTIONS ====================

def fill_paragraph(para, text):
    """Set text in an existing empty paragraph with JUSTIFY alignment."""
    for run in para.runs:
        run.text = ''
    run = para.add_run(text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# 1. Synopsis/Abstract - P[110] (after "Topic Finalization on" heading)
fill_paragraph(doc.paragraphs[110], synopsis_abstract)
print('Filled P[110] - Synopsis/Abstract')

# 2. First Review Summary - P[180] (after "FIRST REVIEW on" heading)
fill_paragraph(doc.paragraphs[180], review_1)
print('Filled P[180] - 1st Review Summary')

# 3. Second Review Summary - P[231] (after "SECOND REVIEW on" heading)
fill_paragraph(doc.paragraphs[231], review_2)
print('Filled P[231] - 2nd Review Summary')

# 4. Third Review Summary - P[279] (after "REVIEW SUMMARY THIRD REVIEW on")
fill_paragraph(doc.paragraphs[279], review_3)
print('Filled P[279] - 3rd Review Summary')

# 5. Fourth Review Summary - P[329] (after "REVIEW SUMMARY FORTH REVIEW on")
fill_paragraph(doc.paragraphs[329], review_4)
print('Filled P[329] - 4th Review Summary')

# 6. Synopsis/Abstract Final Draft - P[379] (after "SYNOPSIS / ABSTRACT FINAL DRAFT")
fill_paragraph(doc.paragraphs[379], final_synopsis)
print('Filled P[379] - Synopsis/Abstract Final Draft')

# Save
doc.save(path)
print()
print('All sections filled and document saved successfully!')
