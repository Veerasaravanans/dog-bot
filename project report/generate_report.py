"""
DogBot Recon System - Project Report Generator
Generates the B.E. project report as .docx matching 22MT060_22MT022 format.
5-Chapter structure: Ch1 Intro | Ch2 Literature | Ch3 Methodology (incl. Hardware & Software)
                     Ch4 Results & Testing | Ch5 Conclusion & Future Scope

Students : VEERA SARAVANAN S (22MT060)  |  JANISHA G (22MT022)
College  : Chennai Institute of Technology (Autonomous)
Dept     : Mechatronics Engineering
Date     : March 2026
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---------------------------------------------------------------------------
# CONSTANTS
# ---------------------------------------------------------------------------
TITLE    = "AUTONOMOUS RECONNAISSANCE WITH AI-POWERED VISION AND PATH PLANNING"
COLLEGE  = "CHENNAI INSTITUTE OF TECHNOLOGY"
COLLEGE2 = "(An Autonomous Institution, Affiliated to Anna University, Chennai)"
DEPT     = "MECHATRONICS ENGINEERING"
DEGREE   = "BACHELOR OF ENGINEERING"
S1_NAME  = "VEERA SARAVANAN S"
S1_REG   = "22MT060"
S2_NAME  = "JANISHA G"
S2_REG   = "22MT022"
HOD_NAME = "Dr. S. CHANDRAVADHANA M.E., Ph.D."
SUP_NAME = "Mr. J. JAI GANESH M.E."
CHAIRMAN = "Shri. P. SRIRAM"
PRINCIPAL= "Dr. A. RAMESH, M.E., Ph.D."
YEAR     = "MARCH 2026"
TNR      = "Times New Roman"

# ---------------------------------------------------------------------------
# HELPERS
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def para(doc, text, size=12, bold=False, italic=False,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, sa=6, sb=0,
         indent=None, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.line_spacing = Pt(22)
    if indent:
        p.paragraph_format.first_line_indent = Cm(indent)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = TNR
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p


def center(doc, text, size=12, bold=False, italic=False, sa=6, sb=0):
    return para(doc, text, size, bold, italic,
                align=WD_ALIGN_PARAGRAPH.CENTER, sa=sa, sb=sb)


def heading_section(doc, num, title, level=1):
    h = doc.add_heading(level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after  = Pt(6)
    run = h.add_run(f"{num}     {title}")
    run.font.name  = TNR
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = True
    run.font.size = Pt(14) if level == 1 else Pt(13) if level == 2 else Pt(12)


def chapter_break(doc, number, title):
    doc.add_page_break()
    h = doc.add_heading(level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run(f"CHAPTER \u2013 {number}")
    r.font.name = TNR; r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0, 0, 0); r.bold = True

    h2 = doc.add_heading(level=1)
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = h2.add_run(title.upper())
    r2.font.name = TNR; r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(0, 0, 0); r2.bold = True
    doc.add_paragraph()


def img_box(doc, fig_num, caption, w=5.5, h=3.0):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F0F0F0")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"\n\n\n[Insert Fig {fig_num} \u2013 {caption}]\n\n\n").font.size = Pt(10)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="6" w:color="999999"/>'
        '  <w:left w:val="single" w:sz="6" w:color="999999"/>'
        '  <w:bottom w:val="single" w:sz="6" w:color="999999"/>'
        '  <w:right w:val="single" w:sz="6" w:color="999999"/>'
        '</w:tcBorders>'
    )
    tcPr.append(borders)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after  = Pt(12)
    r1 = cap.add_run(f"Fig {fig_num} ")
    r1.bold = True; r1.font.size = Pt(10); r1.font.name = TNR
    r2 = cap.add_run(caption)
    r2.font.size = Pt(10); r2.font.name = TNR


def table_with_data(doc, headers, rows, col_widths=None,
                    table_num=None, caption=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True; r.font.size = Pt(10); r.font.name = TNR
        set_cell_shading(cell, "D9E2F3")
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            r.font.size = Pt(10); r.font.name = TNR
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    if caption and table_num:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(4)
        cap.paragraph_format.space_after  = Pt(12)
        r1 = cap.add_run(f"Table {table_num}: ")
        r1.bold = True; r1.font.size = Pt(10); r1.font.name = TNR
        r2 = cap.add_run(caption)
        r2.font.size = Pt(10); r2.font.name = TNR
    return t


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.line_spacing = Pt(20)
        if p.runs:
            p.runs[0].text = item
            p.runs[0].font.size = Pt(12)
            p.runs[0].font.name = TNR
        else:
            r = p.add_run(item)
            r.font.size = Pt(12); r.font.name = TNR


def code_block(doc, code_text):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.cell(0, 0)
    set_cell_shading(cell, "F2F2F2")
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(code_text)
    r.font.size = Pt(9); r.font.name = 'Courier New'
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    doc.add_paragraph()


def remove_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:color="auto"/>'
        '  <w:bottom w:val="none" w:sz="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:color="auto"/>'
        '</w:tcBorders>'
    )
    tcPr.append(tcBorders)


# ---------------------------------------------------------------------------
# FRONT MATTER
# ---------------------------------------------------------------------------

def add_title_page(doc):
    for _ in range(4):
        doc.add_paragraph()

    center(doc, TITLE, 16, bold=True, sa=14)
    doc.add_paragraph()
    center(doc, "A PROJECT REPORT", 13, bold=True, sa=6)
    doc.add_paragraph()
    center(doc, "Submitted by", 12, sa=6)
    doc.add_paragraph()

    # Two students side by side via a no-border table
    t = doc.add_table(rows=2, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    names = [(S1_NAME, S1_REG), (S2_NAME, S2_REG)]
    for ri, (name, reg) in enumerate(names):
        cell_n = t.rows[ri].cells[0]
        cell_r = t.rows[ri].cells[1]
        pn = cell_n.paragraphs[0]
        pr = cell_r.paragraphs[0]
        pn.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rn = pn.add_run(name)
        rr = pr.add_run(reg)
        for r in (rn, rr):
            r.bold = True; r.font.size = Pt(13); r.font.name = TNR
        remove_border(cell_n); remove_border(cell_r)

    doc.add_paragraph()
    center(doc, "in partial fulfillment for the award of the degree", 12, sa=4)
    center(doc, "of", 12, sa=4)
    doc.add_paragraph()
    center(doc, DEGREE, 14, bold=True, sa=4)
    center(doc, "in", 12, sa=4)
    center(doc, DEPT, 14, bold=True, sa=10)
    doc.add_paragraph()
    center(doc, "(Autonomous)", 12, sa=4)
    center(doc, COLLEGE + " (AUTONOMOUS)", 13, bold=True, sa=4)
    center(doc, COLLEGE2, 11, sa=4)
    center(doc, YEAR, 13, bold=True, sa=4)


def add_bonafide_certificate(doc):
    doc.add_page_break()
    for _ in range(2):
        doc.add_paragraph()

    center(doc, COLLEGE, 14, bold=True, sa=4)
    center(doc, COLLEGE2, 11, sa=10)
    center(doc, "BONAFIDE CERTIFICATE", 14, bold=True, sa=12)

    text = (
        f'Certified that this project report entitled \u201c{TITLE}\u201d '
        f'is the bonafide work of {S1_NAME} ({S1_REG}), {S2_NAME} ({S2_REG}) '
        f'who carried out the project work under my supervision.'
    )
    para(doc, text, 12, indent=1.27, sa=14)

    para(doc, 'Certified that the above students have attended the viva voce '
             'examination held on \u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026', 12, indent=1.27, sa=30)

    # Signature block
    t = doc.add_table(rows=3, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = [
        ("INTERNAL EXAMINAR", "EXTERNAL EXAMINAR"),
        ("SIGNATURE", "SIGNATURE"),
        (HOD_NAME + "\nHEAD OF THE DEPARTMENT\nProfessor\nDepartment of Mechatronics Engineering,\nChennai Institute of Technology,\nChennai \u2013 600069.",
         SUP_NAME + "\nSUPERVISOR\nAssistant Professor\nDepartment of Mechatronics Engineering,\nChennai Institute of Technology,\nChennai \u2013 600069."),
    ]
    for ri, (l, r_txt) in enumerate(labels):
        for ci, txt in enumerate((l, r_txt)):
            cell = t.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(txt)
            run.font.size = Pt(11)
            run.font.name  = TNR
            if ri == 0 or ri == 1:
                run.bold = True
            remove_border(cell)


def add_acknowledgement(doc):
    doc.add_page_break()
    for _ in range(2):
        doc.add_paragraph()

    center(doc, "ACKNOWLEDGEMENT", 14, bold=True, sa=12)

    ack_paras = [
        f'We express our gratitude to our Chairman {CHAIRMAN} and all trust members '
        f'of Chennai Institute of Technology for providing the facility and opportunity '
        f'to do this project as a part of our undergraduate course.',

        f'We are grateful to our Principal {PRINCIPAL}, for providing us the facility '
        f'and encouragement during the course of our work.',

        f'We sincerely thank our Head of the Department and Project Guide, {HOD_NAME}, '
        f'Department of Mechatronics Engineering for having provided us valuable guidance, '
        f'resources and timely suggestions throughout our work.',

        'We would like to extend our thanks to our faculty coordinators of the Department '
        'of Mechatronics Engineering, for their valuable suggestions throughout this project.',

        'We wish to extend our sincere thanks to all Faculty members of the Department of '
        'Mechatronics Engineering for their valuable suggestions and their kind cooperation '
        'for the successful completion of our project.',

        'We wish to acknowledge the help received from the Lab Instructors of the Department '
        'of Mechatronics Engineering and others for providing valuable suggestions and for the '
        'successful completion of the project.',
    ]
    for txt in ack_paras:
        para(doc, txt, 12, indent=1.27, sa=10)

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"{S1_NAME}\n{S2_NAME}")
    r.font.size = Pt(12); r.font.name = TNR; r.bold = True


def add_abstract(doc):
    doc.add_page_break()
    center(doc, "ABSTRACT", 14, bold=True, sa=12)

    abstract_text = (
        'The DogBot Recon System is an autonomous reconnaissance robot for indoor and '
        'outdoor navigation using real-time computer vision and AI-driven control. It '
        'integrates an ESP32-CAM for vision and motor control with a Python backend that '
        'handles image processing, obstacle detection, path planning, and autonomous movement. '
        'Multi-sensor fusion combines background subtraction, edge detection, optical flow, '
        'gradient analysis, vertical structure detection, and adaptive floor color detection, '
        'while object recognition and tracking use a lightweight YOLOv8n-seg model. A '
        'low-latency pipeline builds an occupancy grid, predicts obstacle motion with Kalman '
        'filtering, evaluates paths, and outputs steering and speed with emergency stop for '
        'nearby obstacles.\n\n'
        'A dual-layer AI Decision Engine uses a fast local planner with a cloud fallback for '
        'low-confidence cases, supported by multi-frame voting for stability. Remote '
        'connectivity is enabled via secure MQTT, and a FastAPI dashboard with WebSocket '
        'streaming provides live video, telemetry, and manual or semi-autonomous control. '
        'The hardware includes an ESP32-CAM with OV2640 camera, L298N motor driver, four '
        'N20 gear motors, and a battery pack, achieving real-time navigation at about 15 FPS '
        'with low latency.'
    )
    para(doc, abstract_text, 12, indent=1.27, sa=8)


def add_table_of_contents(doc):
    doc.add_page_break()
    center(doc, "TABLE OF CONTENTS", 14, bold=True, sa=12)

    toc = [
        ("",          "ABSTRACT",                                    "iv"),
        ("",          "LIST OF FIGURES",                             "viii"),
        ("",          "LIST OF TABLES",                              "ix"),
        ("1.",        "INTRODUCTION",                                "1"),
        ("1.1",       "BACKGROUND",                                  "1"),
        ("1.2",       "PROBLEM STATEMENT",                           "2"),
        ("1.3",       "OBJECTIVES",                                  "3"),
        ("1.4",       "SCOPE OF THE PROJECT",                        "4"),
        ("2.",        "LITERATURE REVIEW",                           "6"),
        ("2.1",       "INTRODUCTION",                                "6"),
        ("2.2",       "EVOLUTION OF RECONNAISSANCE ROBOTICS",        "6"),
        ("2.3",       "REVIEW OF RELEVANT RESEARCH PAPERS",          "7"),
        ("2.3.1",     "AUTONOMOUS MOBILE ROBOTS",                    "7"),
        ("2.3.2",     "YOLO v8 OBJECT DETECTION",                    "9"),
        ("2.3.3",     "PATH PLANNING & OCCUPANCY GRIDS",             "11"),
        ("2.3.4",     "ESP32 & MQTT",                                "13"),
        ("3.",        "METHODOLOGY",                                  "16"),
        ("3.1",       "SYSTEM OVERVIEW",                             "16"),
        ("3.2",       "DATA FLOW & PROCESSING",                      "18"),
        ("3.3",       "COMMUNICATION PROTOCOL",                      "19"),
        ("3.4",       "CONTROL MODE STRATEGY",                       "20"),
        ("3.5",       "CONFIGURATION MANAGEMENT APPROACH",           "21"),
        ("3.6",       "HARDWARE DESIGN AND IMPLEMENTATION",          "23"),
        ("3.6.1",     "INTRODUCTION",                                "23"),
        ("3.6.2",     "CUSTOM 3D-PRINTED CHASSIS DESIGN",            "23"),
        ("3.6.3",     "ESP32-CAM MODULE",                            "25"),
        ("3.6.4",     "L293D MOTOR DRIVER MODULE",                   "26"),
        ("3.6.5",     "DC MOTORS, WHEELS & LEG SYSTEM",              "28"),
        ("3.6.6",     "POWER SUPPLY",                                "29"),
        ("3.6.7",     "CIRCUIT DIAGRAM & PIN MAPPING",               "30"),
        ("3.6.8",     "HARDWARE ASSEMBLY",                           "31"),
        ("3.7",       "SOFTWARE IMPLEMENTATION",                     "33"),
        ("3.7.1",     "DEVELOPMENT ENVIRONMENT & TOOLS",             "33"),
        ("3.7.2",     "BACKEND FRAMEWORK",                           "34"),
        ("3.7.3",     "ESP32 FIRMWARE",                              "35"),
        ("3.7.4",     "COMPUTER VISION PIPELINE",                    "35"),
        ("3.7.5",     "OBJECT DETECTION",                            "37"),
        ("3.7.6",     "PATH PLANNING ENGINE",                        "38"),
        ("3.7.7",     "AI DECISION ENGINE",                          "39"),
        ("3.7.8",     "MQTT BRIDGE & REMOTE CONNECTIVITY",           "40"),
        ("3.7.9",     "WEB DASHBOARD",                               "41"),
        ("4.",        "RESULTS & TESTING",                           "42"),
        ("4.1",       "TESTING ENVIRONMENT SETUP",                   "42"),
        ("4.2",       "OBSTACLE DETECTION PERFORMANCE",              "43"),
        ("4.3",       "PATH PLANNING PERFORMANCE",                   "44"),
        ("4.4",       "REAL-TIME FPS & LATENCY MEASUREMENTS",        "45"),
        ("4.5",       "MQTT REMOTE CONNECTIVITY TESTING",            "46"),
        ("4.6",       "DASHBOARD DEMONSTRATION",                     "47"),
        ("5.",        "CONCLUSION & FUTURE SCOPE",                   "49"),
        ("5.1",       "CONCLUSIONS",                                 "49"),
        ("5.2",       "FUTURE ENHANCEMENTS",                         "50"),
        ("",          "REFERENCES",                                  "52"),
        ("",          "PO & PSO ATTAINMENT",                         "55"),
    ]

    # Header row
    th = doc.add_table(rows=1, cols=3)
    th.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, txt in enumerate(["CHAPTER NO.", "TITLE", "PAGE NO."]):
        cell = th.rows[0].cells[ci]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        r.bold = True; r.font.size = Pt(11); r.font.name = TNR
        set_cell_shading(cell, "D9E2F3")
    th.rows[0].cells[0].width = Inches(1.2)
    th.rows[0].cells[1].width = Inches(4.5)
    th.rows[0].cells[2].width = Inches(0.8)

    for num, title, pg in toc:
        row = th.add_row()
        for ci, txt in enumerate((num, title, pg)):
            cell = row.cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci != 1 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(txt)
            r.font.size = Pt(10); r.font.name = TNR
            if num.endswith(".") or num in ("", ):
                r.bold = (num.endswith("."))


def add_list_of_figures(doc):
    doc.add_page_break()
    center(doc, "LIST OF FIGURES", 14, bold=True, sa=12)

    figs = [
        ("Fig 1.1.1",  "Dog Bot Recon System Design",                    "5"),
        ("Fig 2.3.2",  "Yolo v8 Architecture",                           "10"),
        ("Fig 2.3.3",  "Occupancy Grid Representation",                  "12"),
        ("Fig 3.1.1",  "System Block Diagram",                           "17"),
        ("Fig 3.2.1",  "Frame Processing Pipeline Flowchart",            "18"),
        ("Fig 3.3.1",  "Communication Architecture",                     "19"),
        ("Fig 3.4.1",  "Manual & Semi-Auto Mode",                        "20"),
        ("Fig 3.6.1",  "Solidworks CAD Model",                           "24"),
        ("Fig 3.6.2",  "3D-Printed Chassis Parts",                       "24"),
        ("Fig 3.6.3",  "ESP32-CAM Module",                               "26"),
        ("Fig 3.6.4",  "L293D Motor Driver",                             "27"),
        ("Fig 3.6.5",  "N20 Motor DC Motor with wheels",                 "28"),
        ("Fig 3.6.6",  "Circuit Diagram",                                "30"),
        ("Fig 3.6.7",  "Motor Direction Truth Table",                    "31"),
        ("Fig 3.6.8",  "Hardware Assembly",                              "32"),
        ("Fig 3.7.1",  "AI Decision Engine Dual-Layer Architecture",     "39"),
        ("Fig 4.1.1",  "Assembled Hardware",                             "42"),
        ("Fig 4.2.1",  "Object Detection Accuracy by Distance range",    "43"),
        ("Fig 4.2.2",  "Object Segmentation & Tracking",                 "44"),
        ("Fig 4.3.1",  "Path Planner Performance Metrics",               "45"),
        ("Fig 4.6.1",  "Dashboard Demonstration",                        "47"),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, txt in enumerate(["FIGURE NO.", "TITLE", "PAGE NO."]):
        cell = t.rows[0].cells[ci]
        cell.text = ""
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt); r.bold = True; r.font.size = Pt(10); r.font.name = TNR
        set_cell_shading(cell, "D9E2F3")
    for fn, title, pg in figs:
        row = t.add_row()
        for ci, txt in enumerate((fn, title, pg)):
            cell = row.cells[ci]; cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci != 1 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(txt); r.font.size = Pt(10); r.font.name = TNR
    t.rows[0].cells[0].width = Inches(1.2)
    t.rows[0].cells[1].width = Inches(4.5)
    t.rows[0].cells[2].width = Inches(0.8)


def add_list_of_tables(doc):
    doc.add_page_break()
    center(doc, "LIST OF TABLES", 14, bold=True, sa=12)

    tbls = [
        ("Table 3.5.1", "Configuration System Sections",          "21"),
        ("Table 3.6.1", "ESP32-CAM Module Specifications",        "25"),
        ("Table 3.6.2", "L293D Motor Driver Specifications",      "27"),
        ("Table 3.7.1", "Software Dependencies & Tools",          "33"),
        ("Table 4.4.1", "System Performance Summary",             "45"),
        ("Table 4.5.1", "MQTT Remote Connectivity Test Results",  "46"),
    ]
    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, txt in enumerate(["TABLE NO.", "TITLE", "PAGE NO."]):
        cell = t.rows[0].cells[ci]; cell.text = ""
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt); r.bold = True; r.font.size = Pt(10); r.font.name = TNR
        set_cell_shading(cell, "D9E2F3")
    for tn, title, pg in tbls:
        row = t.add_row()
        for ci, txt in enumerate((tn, title, pg)):
            cell = row.cells[ci]; cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci != 1 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(txt); r.font.size = Pt(10); r.font.name = TNR


# ---------------------------------------------------------------------------
# CHAPTER 1 - INTRODUCTION
# ---------------------------------------------------------------------------

def add_chapter1(doc):
    chapter_break(doc, "1", "INTRODUCTION")

    heading_section(doc, "1.1", "BACKGROUND")
    para(doc,
         'The field of autonomous mobile robotics has witnessed remarkable growth over the past '
         'decade, driven by advances in embedded computing, computer vision, machine learning, and '
         'wireless communication technologies. Autonomous robots are increasingly being deployed in '
         'applications ranging from warehouse logistics and agricultural monitoring to search and '
         'rescue operations, environmental surveillance, and military reconnaissance. These systems '
         'must perceive their environment, make intelligent decisions, and navigate safely without '
         'continuous human intervention.', 12, indent=1.27)

    para(doc,
         'The emergence of low-cost, high-performance microcontrollers such as the ESP32 series has '
         'democratized robotics development. The ESP32-CAM module, in particular, combines a '
         'dual-core 240 MHz processor, built-in Wi-Fi and Bluetooth, and a 2-megapixel camera in a '
         'compact, affordable package. This makes it possible to build vision-capable robots at a '
         'fraction of the cost of traditional platforms. However, the limited computational power of '
         'such microcontrollers necessitates offloading intensive processing tasks such as deep '
         'learning inference and path planning to an external backend server.', 12, indent=1.27)

    para(doc,
         'Simultaneously, breakthroughs in real-time object detection models like YOLO (You Only '
         'Look Once), now in its eighth generation (YOLOv8), have made it possible to detect, '
         'classify, and segment objects in video streams at high frame rates even on CPUs without '
         'dedicated GPU hardware. When combined with classical computer vision techniques such as '
         'background subtraction, optical flow, and edge detection, these deep learning models '
         'provide robust multi-modal perception capabilities.', 12, indent=1.27)

    para(doc,
         'The motivation behind the Dog Bot Recon System stems from the need for an affordable, '
         'modular, and intelligent reconnaissance robot that can autonomously explore unknown '
         'environments, detect and avoid obstacles in real time, and be controlled remotely from '
         'anywhere in the world via the internet. The project aims to demonstrate that sophisticated '
         'autonomous navigation can be achieved by combining inexpensive hardware with modern '
         'software engineering practices, real-time computer vision, and AI-powered decision making.',
         12, indent=1.27)

    heading_section(doc, "1.2", "PROBLEM STATEMENT")
    para(doc,
         'Existing autonomous robot platforms suffer from several limitations that restrict their '
         'accessibility and practical deployment:', 12, indent=1.27)
    bullets(doc, [
        'High Cost: Commercial autonomous platforms such as Boston Dynamics Spot, Clearpath '
        'Robotics Jackal, or TurtleBot3 cost thousands of dollars, making them inaccessible for '
        'educational institutions and individual researchers.',
        'Limited Remote Access: Most hobby-grade robots operate only on local networks, preventing '
        'remote monitoring and control over the internet from distant locations.',
        'Single-Method Detection: Many low-cost robots rely on a single obstacle detection method '
        '(typically ultrasonic sensors), which provides limited spatial resolution and cannot '
        'identify object types.',
        'No Intelligent Path Planning: Budget robots often use simple reactive behaviors (stop and '
        'turn) rather than predictive path planning that considers obstacle velocities and future '
        'positions.',
        'Fragmented Software: The software for such systems is often monolithic and difficult to '
        'extend, lacking proper service-oriented architecture for modularity.',
    ])
    para(doc,
         'This project builds a modular, cost-effective autonomous robot with vision-based obstacle '
         'detection, deep learning recognition, predictive path planning using Kalman filtering, AI '
         'decision-making, and MQTT-based remote connectivity.', 12, indent=1.27)

    heading_section(doc, "1.3", "OBJECTIVES")
    para(doc, 'The primary objectives of the Dog Bot Recon System project are:', 12, indent=1.27)
    bullets(doc, [
        'To design and build an autonomous reconnaissance robot using the ESP32-CAM '
        'microcontroller as the hardware platform with an L298N motor driver for differential '
        'drive control.',
        'To develop a multi-sensor fusion computer vision pipeline that combines six independent '
        'detection methods (background subtraction, edge detection, gradient magnitude, optical '
        'flow, thin edge detection, and adaptive floor color detection) for robust obstacle '
        'identification.',
        'To integrate YOLOv8n-seg deep learning model for real-time object detection and instance '
        'segmentation with persistent multi-object tracking using the ByteTrack algorithm.',
        'To implement a five-stage real-time path planning engine consisting of occupancy grid '
        'construction, Kalman filter-based obstacle state estimation, cost map generation, '
        'trajectory evaluation with multi-objective scoring, and temporal smoothing for stable '
        'navigation.',
        'To build an AI Decision Engine with dual-layer architecture: fast local path planner as '
        'primary, with cloud-based LLM fallback for complex scenarios, and multi-frame voting for '
        'noise-resistant direction decisions.',
        'To establish worldwide remote connectivity using MQTT protocol over TLS-encrypted '
        'connections, enabling motor control and status monitoring from any internet-connected '
        'device.',
        'To create a real-time web dashboard with live video streaming, telemetry display, manual '
        'and semi-automatic control modes, and automotive-style visualization using FastAPI and '
        'WebSocket protocols.',
        'To achieve real-time performance targets of 15 FPS processing rate, sub-2ms path planning '
        'latency, and sub-150ms end-to-end local latency.',
    ])

    heading_section(doc, "1.4", "SCOPE OF THE PROJECT")
    para(doc,
         'The scope of the Dog Bot Recon System encompasses the complete design, development, and '
         'testing of an autonomous reconnaissance robot platform. The project covers the following '
         'areas:', 12, indent=1.27)
    bullets(doc, [
        'Hardware Scope: The hardware design focuses on integrating the AI-Thinker ESP32-CAM, '
        'L298N motor driver, four N20 DC gear motors, chassis, and 4\u00d7AA battery supply, '
        'including circuit connections and pin mapping.',
        'Software Scope: The software development spans three major areas. First, the ESP32 '
        'firmware written in Arduino C++ handles camera streaming, motor control, WiFi '
        'connectivity, and MQTT communication. Second, the Python backend built with FastAPI '
        'implements the complete vision processing pipeline, machine learning detection, path '
        'planning, AI decision making, and WebSocket-based communication. Third, the web '
        'dashboard provides the user interface for monitoring and control.',
        'Connectivity Scope: Remote connectivity through MQTT protocol with TLS encryption '
        'enables worldwide access. The system supports both local HTTP communication and remote '
        'MQTT communication with automatic failover between modes.',
    ])
    img_box(doc, "1.1.1", "Dog Bot Recon System Design")


# ---------------------------------------------------------------------------
# CHAPTER 2 - LITERATURE SURVEY
# ---------------------------------------------------------------------------

def add_chapter2(doc):
    chapter_break(doc, "2", "LITERATURE SURVEY")

    heading_section(doc, "2.1", "INTRODUCTION")
    para(doc,
         'The field of autonomous robotic reconnaissance has experienced rapid advancement driven '
         'by improvements in sensor technology, artificial intelligence, and wireless communication '
         'systems. Research in this domain spans multiple disciplines including robotics, computer '
         'vision, machine learning, embedded systems, and IoT communication protocols. Understanding '
         'the current state of technology and identifying gaps in existing solutions forms the '
         'foundation for developing an effective reconnaissance robotics system such as the Dog Bot.',
         12, indent=1.27)

    para(doc,
         'Contemporary robotic platforms have demonstrated increasing sophistication in navigation, '
         'environmental adaptation, and autonomous decision-making. The integration of multiple '
         'sensor modalities with advanced processing capabilities has enabled robots to perform '
         'complex tasks in challenging environments previously accessible only to human operators. '
         'In particular, the convergence of lightweight deep-learning models, low-cost '
         'microcontrollers with integrated cameras, and cloud-based large language models has opened '
         'new possibilities for building affordable yet capable autonomous systems.', 12, indent=1.27)

    heading_section(doc, "2.2", "EVOLUTION OF RECONNAISSANCE ROBOTICS")
    para(doc,
         'Early robotic surveillance systems were mainly remote-controlled vehicles with basic '
         'cameras, requiring continuous human monitoring and operating only in controlled '
         'environments. The introduction of autonomous navigation reduced human dependence and '
         'improved operational flexibility.', 12, indent=1.27)

    para(doc,
         'In the early 2000s, probabilistic robotics improved robot reliability through methods '
         'such as Kalman and particle filtering, enabling accurate localization and mapping under '
         'sensor uncertainty. After 2012, deep learning enabled real-time object detection and '
         'shifted surveillance robots toward intelligent perception.', 12, indent=1.27)

    para(doc,
         'Recent advancements in low-cost embedded hardware such as ESP32-CAM have made robotic '
         'surveillance more affordable and accessible. Modern systems now combine machine learning '
         'for object recognition with emerging LLM-based reasoning to support smarter '
         'decision-making, faster responses, and reduced manual supervision.', 12, indent=1.27)

    heading_section(doc, "2.3", "REVIEW OF RELEVANT RESEARCH PAPERS")

    heading_section(doc, "2.3.1", "AUTONOMOUS MOBILE ROBOTS", level=2)

    para(doc, 'Paper 1: \u201cProbabilistic Robotics\u201d by Thrun, S., Burgard, W., and Fox, D. (2005)',
         12, bold=True, indent=1.27)
    para(doc,
         'Thrun, Burgard, and Fox authored the seminal textbook \u201cProbabilistic Robotics,\u201d '
         'published by MIT Press, which established the theoretical foundation for modern autonomous '
         'mobile robots. The book presents probabilistic frameworks including Bayesian filters, '
         'particle filters, and extended Kalman filters for the core problems of localization, '
         'mapping, and Simultaneous Localization and Mapping (SLAM). Their central thesis is that '
         'probabilistic approaches can gracefully handle the inherent uncertainty in sensor '
         'measurements and robot motion, enabling reliable autonomous navigation in real-world '
         'environments where deterministic methods fail.', 12, indent=1.27)

    para(doc,
         'In the context of the DogBot system, the probabilistic reasoning principles from this '
         'work inform the obstacle state estimation module, which maintains belief distributions '
         'over obstacle positions and velocities. The Bayesian update framework from Thrun et al. '
         'underpins the Kalman filter-based obstacle tracker that enables predictive collision '
         'avoidance.', 12, indent=1.27)

    para(doc,
         'Paper 2: \u201cA Formal Basis for the Heuristic Determination of Minimum Cost Paths\u201d '
         'by Hart, P. E., Nilsson, N. J., and Raphael, B. (1968)',
         12, bold=True, indent=1.27)
    para(doc,
         'Hart, Nilsson, and Raphael introduced the A* search algorithm in their landmark paper '
         'published in IEEE Transactions on Systems Science and Cybernetics. A* combines the actual '
         'cost from the start node (g-cost) with a heuristic estimate of the remaining cost to the '
         'goal (h-cost), guaranteeing optimal paths when the heuristic is admissible. This elegant '
         'formulation unified previous graph-search methods into a single, provably optimal '
         'framework with over 12,000 citations.', 12, indent=1.27)

    para(doc,
         'While the DogBot system employs a trajectory evaluation approach rather than explicit '
         'graph search for real-time local navigation, A* provides the theoretical backdrop for '
         'understanding optimal path finding. The cost-based reasoning introduced by Hart et al. '
         '\u2014 evaluating paths by combining traversal cost with goal proximity \u2014 directly '
         'inspires the multi-objective trajectory scoring used in the DogBot\u2019s path planner.',
         12, indent=1.27)

    heading_section(doc, "2.3.2", "YOLO V8 OBJECT DETECTION", level=2)

    para(doc,
         'Paper 1: \u201cYou Only Look Once: Unified, Real-Time Object Detection\u201d '
         'by Redmon, J., Divvala, S., Girshick, R., and Farhadi, A. (2016)',
         12, bold=True, indent=1.27)
    para(doc,
         'Redmon et al. introduced the YOLO framework at CVPR 2016, treating object detection as a '
         'single regression problem rather than a two-stage pipeline. The network divides the input '
         'image into a grid and simultaneously predicts bounding boxes and class probabilities for '
         'all grid cells in one forward pass, achieving 45 FPS on a Titan X GPU with competitive '
         'accuracy. This unified approach enabled real-time detection far beyond the capabilities '
         'of contemporary two-stage detectors like Faster R-CNN.', 12, indent=1.27)

    para(doc,
         'Paper 2: \u201cUltralytics YOLOv8\u201d by Jocher, G., Chaurasia, A., and Qiu, J. (2023)',
         12, bold=True, indent=1.27)
    para(doc,
         'The YOLOv8 architecture from Ultralytics represents the eighth generation of the YOLO '
         'family, incorporating a new backbone with C2f modules (cross-stage partial with two '
         'convolutions), an anchor-free detection head, and native support for segmentation, '
         'pose estimation, and classification tasks alongside detection. The nano variant '
         '(YOLOv8n) achieves 37.3 mAP on COCO with only 3.2M parameters and 8.7 GFLOPs, making '
         'it suitable for CPU inference on embedded devices.', 12, indent=1.27)

    para(doc,
         'The DogBot system uses the YOLOv8n-seg variant for instance segmentation, providing '
         'both bounding boxes and pixel-level segmentation masks for detected objects. Segmentation '
         'masks are used to generate more accurate obstacle footprints in the occupancy grid '
         'compared to using bounding boxes alone, particularly for non-rectangular objects like '
         'chairs and people.', 12, indent=1.27)

    img_box(doc, "2.3.2", "Yolo v8 Architecture")

    heading_section(doc, "2.3.3", "PATH PLANNING & OCCUPANCY GRIDS", level=2)

    para(doc,
         'Paper 1: \u201cOccupancy Grids: A Representation for Dynamic Environments\u201d '
         'by Elfes, A. (1989)',
         12, bold=True, indent=1.27)
    para(doc,
         'Alberto Elfes introduced the occupancy grid representation in his seminal paper in IEEE '
         'Control Systems Magazine (Vol. 9, No. 1, pp. 35\u201343). An occupancy grid divides the '
         'robot\u2019s environment into discrete cells, each storing a probability of being '
         'occupied by an obstacle. The grid is updated using sonar or camera measurements through '
         'a Bayesian update rule, providing a probabilistic map that gracefully handles sensor '
         'noise and uncertainty.', 12, indent=1.27)

    para(doc,
         'The DogBot system implements a 60\u00d780 occupancy grid with 5 cm resolution, covering '
         'a 3m \u00d7 4m area in front of the robot. Each cell stores occupancy probability updated '
         'from four independent detection sources: CV pipeline obstacles, YOLO segmentation masks, '
         'thin vertical structure detection, and floor color anomalies.', 12, indent=1.27)

    img_box(doc, "2.3.3", "Occupancy Grid Representation")

    para(doc,
         'Paper 2: \u201cThe Dynamic Window Approach to Collision Avoidance\u201d '
         'by Fox, D., Burgard, W., and Thrun, S. (1997)',
         12, bold=True, indent=1.27)
    para(doc,
         'Fox, Burgard, and Thrun presented the Dynamic Window Approach (DWA) in IEEE Robotics & '
         'Automation Magazine (Vol. 4, No. 1, pp. 23\u201333). DWA samples feasible velocities '
         'from a dynamic window constrained by the robot\u2019s acceleration limits, evaluating '
         'each velocity pair against an objective function considering obstacle clearance, heading '
         'toward the goal, and forward velocity. This enables real-time local navigation in dynamic '
         'environments where obstacles may move or appear unexpectedly.', 12, indent=1.27)

    para(doc,
         'The DogBot system implements a trajectory evaluation approach directly inspired by DWA, '
         'evaluating 41 candidate circular arc trajectories per frame. Each trajectory is scored by '
         'a weighted cost function considering obstacle proximity (weight 3.0), heading deviation '
         '(0.8), smoothness (0.5), clearance (2.0), and forward progress (1.0). The sub-2ms '
         'execution time ensures the planner runs every frame without blocking the processing loop.',
         12, indent=1.27)

    heading_section(doc, "2.3.4", "ESP32 AND MQTT", level=2)

    para(doc,
         'Paper 1: \u201cDesign and Implementation of ESP32-Based IoT Devices\u201d '
         'by Hercog, D., Lerher, T., Truntie, I., and Tezak, O. (2023)',
         12, bold=True, indent=1.27)
    para(doc,
         'Hercog et al. presented their study in the MDPI journal Sensors (Vol. 23, No. 15, '
         'Article 6739), providing a comprehensive investigation of designing IoT devices using '
         'the ESP32 microcontroller family. The authors evaluate the ESP32\u2019s dual-core '
         'architecture for concurrent task execution \u2014 one core handling communication '
         'while the other processes sensor data \u2014 and demonstrate an optimal balance of '
         'processing power, connectivity, and energy efficiency.', 12, indent=1.27)

    para(doc,
         'The DogBot system uses the ESP32-CAM module with dual-core Xtensa LX6 processors at '
         '240 MHz, 520 KB SRAM, 4 MB flash, 4 MB PSRAM, and an OV2640 camera sensor. The '
         'firmware follows the dual-task architecture with one task managing the MJPEG video '
         'stream on port 81 and another handling motor control and MQTT communication on port 80.',
         12, indent=1.27)

    para(doc,
         'Paper 2: \u201cThe Use of MQTT in M2M and IoT Systems: A Survey\u201d '
         'by Mishra, B. and Kertesz, A. (2020)',
         12, bold=True, indent=1.27)
    para(doc,
         'Mishra and Kertesz published this survey in IEEE Access (Vol. 8, pp. 201071\u2013201086), '
         'analysing MQTT usage in IoT systems over two decades. Their findings demonstrate MQTT\u2019s '
         'dominance as a lightweight messaging protocol due to its minimal overhead (as low as '
         '2 bytes per header), publish-subscribe architecture, three QoS levels, and native TLS '
         'support. The survey identifies key challenges including broker scalability and security '
         'in poorly configured deployments.', 12, indent=1.27)

    para(doc,
         'The DogBot system uses HiveMQ Cloud as the MQTT broker with TLS encryption on port 8883, '
         'directly applying the architecture described by Mishra and Kertesz. The ESP32 publishes '
         'status and heartbeat messages and subscribes to motor command topics, enabling remote '
         'control from any location worldwide without requiring direct network connectivity.',
         12, indent=1.27)


# ---------------------------------------------------------------------------
# CHAPTER 3 - METHODOLOGY  (includes Hardware Design & Software Implementation)
# ---------------------------------------------------------------------------

def add_chapter3(doc):
    chapter_break(doc, "3", "METHODOLOGY")

    # ------------------------------------------------------------------
    # 3.1 - 3.5  (original methodology sections)
    # ------------------------------------------------------------------
    heading_section(doc, "3.1", "SYSTEM OVERVIEW")
    para(doc,
         'The Dog Bot Recon System follows a distributed architecture where the ESP32-CAM '
         'microcontroller serves as the hardware interface (camera and motor control), and a '
         'Python-based backend server handles all computationally intensive processing including '
         'computer vision, machine learning, path planning, and AI decision making. A web-based '
         'dashboard provides the human interface for monitoring and control.', 12, indent=1.27)

    para(doc, 'The system comprises three main tiers:', 12, indent=1.27)
    bullets(doc, [
        'Hardware Tier (ESP32-CAM): Camera MJPEG streaming on port 81, motor control HTTP API '
        'on port 80, MQTT client for remote connectivity, WiFi station for network access.',
        'Processing Tier (Python Backend): FastAPI web server, computer vision pipeline '
        '(OpenCV), ML detection (YOLOv8), path planner engine, AI decision engine, MQTT bridge, '
        'WebSocket streaming, and REST API.',
        'Presentation Tier (Web Dashboard): Real-time video display with overlays, telemetry '
        'panels, manual D-pad controls with speed slider, control mode switching, decision log, '
        'and system status indicators.',
    ])

    img_box(doc, "3.1.1", "System Block Diagram")

    para(doc,
         'The ESP32-CAM captures video frames as an MJPEG stream and transmits them over WiFi to '
         'the backend server. The server processes each frame through a multi-stage pipeline: first '
         'the CV pipeline extracts obstacles using six detection methods, then the ML detector '
         'identifies and classifies objects using YOLOv8, followed by the path planner which '
         'computes optimal trajectories, and finally the AI decision engine determines the motor '
         'command. The processed frame with annotations is streamed to the dashboard via WebSocket, '
         'while motor commands are sent back to the ESP32 via MQTT or HTTP.', 12, indent=1.27)

    heading_section(doc, "3.2", "DATA FLOW AND PROCESSING")
    para(doc,
         'The frame processing pipeline is the heart of the system. The FrameManager orchestrates '
         'the following processing steps for every frame at a target rate of 15 FPS:', 12, indent=1.27)

    img_box(doc, "3.2.1", "Frame Processing Pipeline Flowchart")

    para(doc,
         'The pipeline generates a FrameResult object for each processed frame, which contains '
         'the annotated frame as a base64-encoded JPEG, a list of detected obstacles with their '
         'properties (distance, velocity, threat level), ML detections with bounding boxes and '
         'masks, lane status for all three lanes, and system telemetry including FPS, latency, '
         'and connection status. This FrameResult is broadcast to all connected WebSocket '
         'subscribers.', 12, indent=1.27)

    heading_section(doc, "3.3", "COMMUNICATION PROTOCOL")
    para(doc,
         'The system uses three communication protocols, each optimized for its specific purpose:',
         12, indent=1.27)

    para(doc, 'HTTP Protocol:', 12, bold=True, sa=4)
    para(doc,
         'HTTP is used for two purposes: the ESP32 provides an MJPEG stream on port 81 that the '
         'backend consumes for video frames, and a REST API on port 80 for motor control commands '
         '(e.g., GET /motor?dir=forward&speed=200). On the backend side, FastAPI serves a REST API '
         'for health checks, telemetry, and mode switching. HTTP is used for local network '
         'communication where latency is low.', 12, indent=1.27)

    para(doc, 'WebSocket Protocol:', 12, bold=True, sa=4)
    para(doc,
         'WebSocket provides persistent bidirectional connections between the backend and the web '
         'dashboard. Two WebSocket endpoints are used: /ws/video for streaming annotated frames at '
         '15 FPS (FrameResult JSON with base64 JPEG), and /ws/control for real-time control '
         'commands (motor directions, mode changes) and receiving AI decisions and detection alerts.',
         12, indent=1.27)

    img_box(doc, "3.3.1", "Communication Architecture")

    para(doc, 'MQTT Protocol:', 12, bold=True, sa=4)
    para(doc,
         'MQTT enables remote connectivity over the internet. The system uses HiveMQ Cloud as '
         'the broker with TLS encryption on port 8883. Three topics are used: dogbot/cmd/motor '
         'for receiving motor commands (JSON format: {"dir":"forward","speed":200}), '
         'dogbot/status for publishing system status (RSSI, uptime, motor state, IP), and '
         'dogbot/heartbeat for connection liveness monitoring with a 15-second timeout.',
         12, indent=1.27)

    heading_section(doc, "3.4", "CONTROL MODE STRATEGY")
    para(doc,
         'The DogBot system supports two control modes that determine how navigation decisions '
         'are made:', 12, indent=1.27)

    para(doc, 'MANUAL Mode:', 12, bold=True, sa=4)
    para(doc,
         'In manual mode, the human operator has full control of the robot through the web '
         'dashboard D-pad controls or remote MQTT commands. The AI decision engine is disabled, '
         'and all motor commands come directly from user input. The vision pipeline and ML '
         'detection continue to run for monitoring purposes (the operator sees obstacle overlays '
         'and detection boxes) but do not generate motor commands.', 12, indent=1.27)

    img_box(doc, "3.4.1", "Manual and Semi-Auto Mode")

    para(doc, 'SEMI \u2013 AUTO Mode:', 12, bold=True, sa=4)
    para(doc,
         'In semi-automatic mode, the AI decision engine actively controls the robot based on '
         'the path planner output. However, when the operator sends a manual command, AI control '
         'is temporarily paused for 5 seconds to allow the human to override. After the pause '
         'expires, AI control resumes automatically. This mode provides a balance between '
         'autonomous operation and human oversight, allowing intervention when needed.',
         12, indent=1.27)

    heading_section(doc, "3.5", "CONFIGURATION MANAGEMENT APPROACH")
    para(doc,
         'The system uses a centralized configuration management approach based on Pydantic '
         'Settings, which loads configuration from a .env file and provides type validation and '
         'documentation for all settings. The configuration is organized into logical sections:',
         12, indent=1.27)

    table_with_data(doc,
        headers=["Section", "Key Parameters", "Description"],
        rows=[
            ("ESP32",       "stream_url, control_url, timeout",
             "ESP32-CAM connection settings (IP, ports, timeout)"),
            ("VIO LLM",     "base_url, username, api_token, model",
             "Cloud LLM API credentials and model selection"),
            ("MQTT",        "broker_host, port, username, password, use_tls",
             "HiveMQ Cloud broker connection settings"),
            ("CV Pipeline", "contour_min_area, canny thresholds, blur kernel",
             "Computer vision tuning parameters"),
            ("ML Detection","model_path, confidence_threshold, detect_every_n",
             "YOLOv8 model and inference settings"),
            ("Path Planner","grid_width_m, grid_depth_m, cell_size_m",
             "Occupancy grid dimensions and resolution"),
            ("App",         "host, port",
             "FastAPI server binding configuration"),
        ],
        col_widths=[1.2, 2.2, 2.8],
        table_num="3.5.1",
        caption="Configuration System Sections"
    )

    para(doc,
         'An example .env file (.env.example) is provided with the project as a template. The '
         'configuration system ensures that sensitive credentials (API tokens, MQTT passwords) are '
         'never hardcoded in source code and are excluded from version control through .gitignore.',
         12, indent=1.27)

    # ------------------------------------------------------------------
    # 3.6  HARDWARE DESIGN AND IMPLEMENTATION
    # ------------------------------------------------------------------
    heading_section(doc, "3.6", "HARDWARE DESIGN AND IMPLEMENTATION")

    heading_section(doc, "3.6.1", "INTRODUCTION", level=2)
    para(doc,
         'The hardware of the Dog Bot Recon System is designed to form a compact, stable, and '
         'efficient quadruped robotic platform. The system integrates mechanical structure, embedded '
         'control, motor actuation, and power management into a unified architecture.', 12, indent=1.27)
    para(doc,
         'The design emphasizes lightweight construction, structural stability, modular integration, '
         'and reliable operation for indoor and uneven terrain movement. Proper component arrangement '
         'and wiring layout ensure balanced weight distribution and stable performance during '
         'operation.', 12, indent=1.27)

    heading_section(doc, "3.6.2", "CUSTOM 3D-PRINTED CHASSIS DESIGN", level=2)
    para(doc,
         'The Dog Bot chassis is a custom-designed, 3D-printed dog-shaped body created using '
         'Solidworks CAD software. The design adopts a quadruped dog-like form factor with a '
         'distinct head section that houses the ESP32-CAM module, providing a forward-facing '
         'camera window for vision-based navigation. The body measures approximately 150mm \u00d7 '
         '95mm and is printed in black PLA filament using a Fused Deposition Modeling (FDM) '
         '3D printer.', 12, indent=1.27)

    para(doc,
         'The chassis design incorporates several key features: an internal cavity for the L293D '
         'motor driver module and battery cells, mounting points for the four N20 motors, and an '
         'articulated leg system. Each leg consists of multiple 3D-printed pieces \u2014 an upper '
         'leg link and a lower leg link \u2014 assembled together with M2/M3 screws to create '
         'articulated joints. Custom coil springs, also designed in Fusion 360, are placed between '
         'the body and the legs to provide shock absorption during movement over uneven surfaces.',
         12, indent=1.27)

    img_box(doc, "3.6.1", "Solidworks CAD Model")

    para(doc,
         'The 3D-printing process used the following parameters: PLA filament, 0.2mm layer height, '
         '20\u201330% infill density for structural parts, and support structures for overhanging '
         'features such as the head section and leg mounting brackets. The total print time for all '
         'body and leg components was approximately 8\u201312 hours across multiple print sessions.',
         12, indent=1.27)

    img_box(doc, "3.6.2", "3D-Printed Chassis Parts")

    heading_section(doc, "3.6.3", "ESP32-CAM MODULE (AI-THINKER)", level=2)
    para(doc,
         'The AI-Thinker ESP32-CAM is a compact, low-cost development board that integrates an '
         'ESP32-S microprocessor, an OV2640 camera sensor, and supporting peripherals. The module '
         'measures approximately 27mm \u00d7 40.5mm \u00d7 4.5mm and provides a powerful platform '
         'for vision-based IoT applications.', 12, indent=1.27)

    table_with_data(doc,
        headers=["Parameter", "Specification"],
        rows=[
            ("Microprocessor",    "ESP32-S (Dual-core Xtensa LX6, 240 MHz)"),
            ("Camera Sensor",     "OV2640, 2 Megapixel, UXGA (1600\u00d71200)"),
            ("WiFi",              "IEEE 802.11 b/g/n, 2.4 GHz"),
            ("Bluetooth",         "Bluetooth 4.2 BR/EDR + BLE"),
            ("Flash Memory",      "4 MB"),
            ("PSRAM",             "4 MB (external SPI RAM)"),
            ("SRAM",              "520 KB"),
            ("Operating Voltage", "5V (via VCC pin or USB)"),
            ("GPIO Pins",         "16 pins (limited due to camera interface usage)"),
            ("SD Card Slot",      "Micro-SD (shares GPIO with some pins)"),
            ("LED Flash",         "Built-in white LED on GPIO 4"),
            ("Antenna",           "Onboard PCB antenna (optional external via IPEX)"),
            ("Dimensions",        "27mm \u00d7 40.5mm \u00d7 4.5mm"),
        ],
        col_widths=[2.2, 4.0],
        table_num="3.6.1",
        caption="ESP32-CAM Module Specifications"
    )

    img_box(doc, "3.6.3", "ESP32-CAM Module")

    para(doc,
         'The ESP32-CAM provides six GPIO pins for motor control in the DogBot system: GPIO 12, '
         '13, 14, and 15 for motor direction control (connected to L293D IN1\u2013IN4), and '
         'GPIO 2 and GPIO 4 for PWM speed control (connected to L293D EN1 and EN2). Note that '
         'GPIO 4 is shared with the onboard LED flash, which is disabled in the firmware to '
         'prevent interference with PWM output.', 12, indent=1.27)

    heading_section(doc, "3.6.4", "L293D MOTOR DRIVER MODULE", level=2)
    para(doc,
         'The L293D is a quadruple half-H driver IC designed to drive up to four DC motors or '
         'two stepper motors. The module version used in the DogBot system is a 4-channel DC '
         'motor driver module featuring the L293D IC on a blue PCB with screw terminals for '
         'motor and power connections (A\u2212, A+, B\u2212, B+, GND, VIN) and a header pin '
         'row for logic inputs (EN2, IN1, IN3, IN2, IN1, GND, VCC) with removable jumpers on '
         'the enable pins for PWM speed control.', 12, indent=1.27)

    table_with_data(doc,
        headers=["Parameter", "Specification"],
        rows=[
            ("IC",                         "L293D Quadruple Half-H Driver"),
            ("Number of Channels",         "4 (configured as 2 full bridges for Motor A and Motor B)"),
            ("Logic Input Voltage",        "5V (from ESP32 GPIO)"),
            ("Motor Supply Voltage (VIN)", "4.5V \u2013 36V"),
            ("Continuous Current / Channel","600 mA"),
            ("Peak Current / Channel",     "1.2 A"),
            ("Output Clamp Diodes",        "Built-in (internal protection diodes)"),
            ("PWM Speed Control",          "Via EN1 and EN2 pins (remove jumpers for external PWM)"),
            ("PWM Frequency",              "1 kHz (configured in firmware)"),
            ("PWM Resolution",             "8-bit (0\u2013255 speed levels)"),
            ("Module PCB Color",           "Blue"),
            ("Screw Terminals",            "A\u2212, A+, B\u2212, B+, GND, VIN"),
            ("Input Header",               "EN2, IN1, IN3, IN2, IN1, GND, VCC"),
        ],
        col_widths=[2.5, 3.7],
        table_num="3.6.2",
        caption="L293D Motor Driver Specifications"
    )

    img_box(doc, "3.6.4", "L293D Motor Driver")

    heading_section(doc, "3.6.5", "DC MOTORS, WHEELS, AND LEG SYSTEM", level=2)
    para(doc,
         'The DogBot uses four N20 micro DC gear motors rated at 60 RPM operating at 3\u20136V. '
         'The N20 motors are compact (12mm \u00d7 10mm \u00d7 24mm), lightweight gear motors '
         'with a metal gearbox that provides a good balance of speed and torque for indoor '
         'navigation. The system supplies 7.4V from the Li-ion battery pack through the L293D '
         'driver, which regulates the voltage delivered to the motors via PWM duty cycle control.',
         12, indent=1.27)

    para(doc,
         'Each motor is fitted with a rubber tire wheel with a plastic hub, measuring 23mm in '
         'radius (approximately 46mm diameter). The rubber tires provide adequate traction on '
         'indoor surfaces such as tile, wood, and carpet. The wheels are press-fitted onto the '
         'N20 motor shafts.', 12, indent=1.27)

    img_box(doc, "3.6.5", "N20 Micro DC Motor with wheels")

    para(doc,
         'The four motors are configured as two groups: the left pair (front-left and rear-left) '
         'is controlled as Motor A through L293D screw terminals A\u2212 and A+, while the right '
         'pair (front-right and rear-right) is controlled as Motor B through terminals B\u2212 '
         'and B+. This differential drive configuration enables the robot to move forward, '
         'backward, and perform tank-style turns (left and right wheels rotating in opposite '
         'directions).', 12, indent=1.27)

    para(doc,
         'The articulated leg system connects each motor to the chassis body through custom '
         '3D-printed upper and lower leg links. The leg joints are secured with M2/M3 screws '
         'allowing controlled articulation. Custom coil springs installed between the body and '
         'each leg assembly provide suspension, absorbing shocks from uneven terrain and improving '
         'traction by keeping all four wheels in contact with the ground.', 12, indent=1.27)

    heading_section(doc, "3.6.6", "POWER SUPPLY", level=2)
    para(doc,
         'The DogBot system uses a dual power supply approach to ensure clean operation:',
         12, indent=1.27)
    bullets(doc, [
        'Motor Power (Battery): Two 3.7V lithium-ion cells connected in series provide 7.4V for '
        'the motors. This connects to the L293D module VIN screw terminal (the module accepts '
        '4.5\u201336V). The higher voltage from Li-ion cells provides better motor performance '
        'and longer runtime.',
        'Logic Power (USB): The ESP32-CAM is powered via its micro-USB connector from a USB '
        'power bank or computer USB port. This provides a stable 5V supply for the '
        'microprocessor, camera, WiFi, and MQTT operations. Separating motor power from logic '
        'power prevents voltage dips from motor current draw affecting the ESP32 stability.',
        'Ground Connection: A common ground wire connects the ESP32-CAM GND pin to the L293D '
        'GND terminal, ensuring both power domains share the same reference voltage. This is '
        'critical for proper logic-level communication between the ESP32 GPIO outputs and the '
        'L293D logic inputs.',
    ])

    heading_section(doc, "3.6.7", "CIRCUIT DIAGRAM AND PIN MAPPING", level=2)
    para(doc,
         'The complete wiring between the ESP32-CAM and L293D motor driver follows the pin '
         'mapping below. Each connection is color-coded for easy identification during assembly '
         'and debugging. The pin diagram is also available as an interactive SVG visualization '
         'accessible at http://localhost:8000/pin-diagram when the system is running.',
         12, indent=1.27)

    img_box(doc, "3.6.6", "Circuit Diagram")

    para(doc,
         'The motor direction truth table below shows the logic states required on the L293D '
         'input pins to achieve each movement direction. The ESP32 firmware sets these GPIO '
         'outputs accordingly when a motor command is received:', 12, indent=1.27)

    img_box(doc, "3.6.7", "Motor Direction Truth Table")

    heading_section(doc, "3.6.8", "HARDWARE ASSEMBLY", level=2)
    para(doc,
         'The hardware assembly of the DogBot system involves integrating the 3D-printed chassis '
         'with the locomotion and control components in a structured manner. The articulated leg '
         'mechanism is first assembled using mechanical fasteners, and coil springs are incorporated '
         'to provide suspension and stability during movement. Four N20 DC gear motors (60 RPM) '
         'are then mounted onto the chassis and connected to rubber-tired wheels to enable smooth '
         'locomotion. The L293D motor driver module is securely fixed inside the chassis, and the '
         'motors are wired in a left-right parallel configuration to support differential drive '
         'control.', 12, indent=1.27)

    para(doc,
         'The ESP32-CAM module is positioned at the front head section of the chassis to ensure '
         'an optimal camera field of view for navigation and monitoring. A 7.4V Li-ion battery '
         'pack supplies power to both the motors and control electronics. The GPIO pins of the '
         'ESP32-CAM are connected to the L293D motor driver to provide direction control and '
         'PWM-based speed regulation. After completing all wiring connections, firmware is '
         'uploaded to the ESP32-CAM, and the system is validated through motor movement tests '
         'and live video streaming to ensure proper functionality.',
         12, indent=1.27)

    img_box(doc, "3.6.8", "Hardware Assembly")

    # ------------------------------------------------------------------
    # 3.7  SOFTWARE IMPLEMENTATION
    # ------------------------------------------------------------------
    heading_section(doc, "3.7", "SOFTWARE IMPLEMENTATION")

    heading_section(doc, "3.7.1", "DEVELOPMENT ENVIRONMENT AND TOOLS", level=2)
    para(doc,
         'The DogBot Recon System software is developed using Python 3.10+ for the backend server '
         'and Arduino C++ for the ESP32 firmware. The development environment and tools used are:',
         12, indent=1.27)

    table_with_data(doc,
        headers=["Tool / Technology", "Version", "Purpose"],
        rows=[
            ("Python",          "3.10+",    "Backend server language"),
            ("FastAPI",         "\u22650.104.0", "Async web framework with auto-docs"),
            ("Uvicorn",         "\u22650.24.0", "ASGI server for FastAPI"),
            ("OpenCV (cv2)",    "\u22654.8.0", "Computer vision processing"),
            ("Ultralytics",     "\u22658.0.0", "YOLOv8 object detection framework"),
            ("NumPy",           "\u22651.24.0", "Numerical computing for arrays"),
            ("aiomqtt",         "\u22652.0.0", "Async MQTT client library"),
            ("HTTPX",           "\u22650.25.0", "Async HTTP client for ESP32 API"),
            ("Pydantic",        "\u22652.5.0", "Data validation and settings"),
            ("Pillow",          "\u226510.0.0", "Image processing utilities"),
            ("Arduino IDE",     "2.x",      "ESP32 firmware development"),
            ("ESP32 Board Support","2.x",   "Arduino core for ESP32"),
        ],
        col_widths=[2.0, 1.2, 3.0],
        table_num="3.7.1",
        caption="Software Dependencies and Tools"
    )

    heading_section(doc, "3.7.2", "BACKEND FRAMEWORK (FASTAPI)", level=2)
    para(doc,
         'FastAPI is a modern, high-performance Python web framework for building APIs with '
         'automatic OpenAPI documentation, async/await support, and type-based validation via '
         'Pydantic. The DogBot backend uses FastAPI for three purposes: serving the static web '
         'dashboard files, providing REST API endpoints for system monitoring, and managing '
         'WebSocket connections for real-time streaming.', 12, indent=1.27)

    para(doc,
         'The application startup and shutdown are managed through FastAPI\u2019s lifespan context '
         'manager, which ensures all services are started in the correct dependency order '
         '(MQTTBridge \u2192 ESP32Client \u2192 MLDetector \u2192 AIDecisionEngine \u2192 '
         'FrameManager) during startup and stopped in reverse order during shutdown. This pattern '
         'ensures that dependent services are available before their consumers start.',
         12, indent=1.27)

    code_block(doc,
        '# Service startup order in backend/main.py\n'
        'async def lifespan(app: FastAPI):\n'
        '    # Start services in dependency order\n'
        '    await mqtt_bridge.start()\n'
        '    await esp32_client.start()\n'
        '    await ml_detector.start()\n'
        '    await ai_decision.start()\n'
        '    await frame_manager.start()\n'
        '    yield\n'
        '    # Stop in reverse order\n'
        '    await frame_manager.stop()\n'
        '    await ai_decision.stop()\n'
        '    await ml_detector.stop()\n'
        '    await esp32_client.stop()\n'
        '    await mqtt_bridge.stop()'
    )

    heading_section(doc, "3.7.3", "ESP32 FIRMWARE (ARDUINO C++)", level=2)
    para(doc,
         'The ESP32-CAM firmware is written in Arduino C++ and follows a dual-server structure:',
         12, indent=1.27)
    bullets(doc, [
        'HTTP server (port 80) for motor control and status requests',
        'MJPEG stream server (port 81) for live video feed',
    ])
    para(doc,
         'It also connects to HiveMQ Cloud via MQTT (TLS port 8883) for remote motor commands. '
         'The firmware configures camera settings, WiFi connection, motor GPIO pins, PWM speed '
         'control, and implements a failsafe stop mechanism if no command is received within '
         '500 ms.', 12, indent=1.27)

    heading_section(doc, "3.7.4", "COMPUTER VISION PIPELINE (OPENCV)", level=2)
    para(doc,
         'The Computer Vision Pipeline (CV Pipeline) is a major component of the Dog Bot Recon '
         'System, implemented using Python and OpenCV. It processes each incoming camera frame to '
         'generate real-time environmental understanding such as obstacle detection, distance '
         'estimation, motion tracking, threat assessment, and lane-based navigation status.',
         12, indent=1.27)

    para(doc,
         'To improve obstacle detection reliability, the pipeline uses multiple OpenCV-based '
         'methods including background subtraction (MOG2), edge detection, gradient analysis '
         '(Sobel), dense optical flow (Farneback), vertical edge detection, and floor color '
         'anomaly detection. The results from these methods are fused using an OR-based strategy, '
         'meaning any region detected as an obstacle by at least one method is treated as a valid '
         'obstacle, improving detection recall for safety-critical navigation.', 12, indent=1.27)

    para(doc,
         'Obstacle distance is estimated using monocular perspective geometry, where objects '
         'closer to the robot appear lower in the image. The distance is computed using the '
         'formula: distance = DIST_SCALE / (obstacle_y \u2212 vanish_y), where DIST_SCALE is a '
         'calibration constant and vanish_y is the estimated vanishing point. The vanishing point '
         'is updated periodically using Hough line-based estimation.', 12, indent=1.27)

    para(doc,
         'Detected obstacles are tracked using a centroid-based multi-object tracking approach. '
         'Each obstacle is assigned a unique ID, and tracking is maintained using centroid '
         'proximity matching. Using the tracked positions, the system estimates obstacle velocity '
         'and calculates time-to-collision (TTC). Based on distance, motion, and TTC, a threat '
         'level is assigned to each obstacle to support safe decision-making.', 12, indent=1.27)

    para(doc,
         'For navigation guidance, the frame is divided into three lanes (left, center, right), '
         'and each lane is classified as clear, caution, or blocked depending on the closest '
         'detected obstacle. The lane with maximum clearance is recommended as the best path for '
         'movement.', 12, indent=1.27)

    heading_section(doc, "3.7.5", "OBJECT DETECTION (YOLOV8N-SEG)", level=2)
    para(doc,
         'The MLDetector service integrates the YOLOv8n-seg model from Ultralytics for real-time '
         'object detection and instance segmentation. Unlike the CV pipeline which detects generic '
         'obstacles based on visual features, the ML detector identifies and classifies specific '
         'object categories (persons, chairs, tables, bottles, etc.) with semantic labels.',
         12, indent=1.27)

    para(doc, 'Key implementation details:', 12, indent=1.27)
    bullets(doc, [
        'Model: YOLOv8n-seg (nano segmentation) with 3.4M parameters, optimized for CPU inference.',
        'Inference Frequency: Every 3 frames (configurable) to balance accuracy and performance. '
        'Between inference frames, the previous detections with tracked positions are reused.',
        'Tracking: Built-in ByteTrack algorithm provides persistent object IDs across frames, '
        'enabling temporal reasoning about detected objects.',
        'Output Format: Each detection includes class name, confidence score (0\u20131), bounding '
        'box (x, y, w, h), segmentation mask (polygon points), zone classification '
        '(DANGER/CAUTION/SAFE), and track ID for persistence.',
        'Thread Execution: ML inference runs in a thread pool executor to avoid blocking the main '
        'asyncio event loop, ensuring the frame processing pipeline remains responsive.',
    ])

    heading_section(doc, "3.7.6", "PATH PLANNER ENGINE", level=2)
    para(doc,
         'The Path Planner Engine is the main navigation module of the DogBot system. It converts '
         'obstacle and environment data into real-time steering and speed commands, enabling smooth '
         'and predictive robot movement with an average processing latency of approximately 2 ms '
         'per frame.', 12, indent=1.27)

    para(doc,
         'The planner follows a multi-stage pipeline consisting of occupancy grid generation, '
         'obstacle motion estimation, cost-map formation, trajectory evaluation, and temporal '
         'smoothing. A 60\u00d780 occupancy grid is generated with 5 cm resolution, covering a '
         '3m \u00d7 4m area in front of the robot. Each grid cell stores an occupancy probability '
         'ranging from 0.0 (free) to 1.0 (occupied). The grid is updated using inputs from the '
         'CV obstacle pipeline, YOLO detections, thin vertical structure detection, and floor '
         'anomaly detection. A temporal decay factor is applied to reduce the effect of outdated '
         'obstacles and suppress false positives.', 12, indent=1.27)

    para(doc,
         'To improve obstacle tracking stability, the planner applies a Kalman filter for each '
         'detected obstacle using a constant-velocity model. The state vector is defined as '
         '[x, y, vx, vy], allowing estimation of both position and velocity. The Kalman '
         'predict-update cycle generates smooth motion estimates and enables prediction of obstacle '
         'positions for short future time horizons. Using these estimates, the system produces a '
         'weighted cost map and evaluates multiple candidate trajectories, selecting the safest '
         'and most efficient path. Final motion outputs are stabilized using temporal smoothing to '
         'prevent sudden steering changes.', 12, indent=1.27)

    heading_section(doc, "3.7.7", "AI DECISION ENGINE", level=2)
    para(doc,
         'The AI Decision Engine implements a dual-layer decision architecture that combines fast '
         'local planning with intelligent cloud-based reasoning:', 12, indent=1.27)

    img_box(doc, "3.7.1", "AI Decision Engine Dual-Layer Architecture")

    para(doc,
         'The multi-frame voting mechanism accumulates direction votes over a sliding window of '
         '5 frames. A direction change is only issued when one direction receives at least 60% '
         'of votes (3 out of 5 frames). This prevents single-frame noise or transient detections '
         'from causing unnecessary direction changes. The STOP command is exempt from voting and '
         'is always executed immediately for safety.', 12, indent=1.27)

    heading_section(doc, "3.7.8", "MQTT BRIDGE AND REMOTE CONNECTIVITY", level=2)
    para(doc,
         'The MQTT Bridge service provides remote connectivity by connecting to a HiveMQ Cloud '
         'broker over TLS-encrypted MQTT. This enables motor control from any internet-connected '
         'device worldwide. The implementation uses the aiomqtt library (version 2.0.0+) which '
         'provides native asyncio support.', 12, indent=1.27)

    para(doc, 'Three MQTT topics are used:', 12, indent=1.27)
    bullets(doc, [
        'dogbot/cmd/motor (Subscribe): Receives motor commands as JSON messages with direction '
        'and speed fields (e.g., {"dir":"forward","speed":200}).',
        'dogbot/status (Publish): Publishes system status every 5 seconds including RSSI signal '
        'strength, uptime, current motor state, IP address, and motor driver information.',
        'dogbot/heartbeat (Publish): Publishes heartbeat messages for connection liveness '
        'monitoring. If no heartbeat is received within 15 seconds, the connection is considered '
        'lost.',
        'Windows Compatibility: On Windows, a special event loop configuration is required for '
        'MQTT compatibility. The Windows ProactorEventLoop does not support the socket '
        'add_reader/add_writer operations required by the underlying paho-mqtt library. The '
        'run.py startup script addresses this by explicitly creating a SelectorEventLoop before '
        'starting the uvicorn server.',
    ])

    heading_section(doc, "3.7.9", "WEB DASHBOARD (FRONTEND)", level=2)
    para(doc,
         'The web dashboard is a single-page application served as static HTML/CSS/JavaScript '
         'by the FastAPI backend. It connects to the backend via two WebSocket connections '
         '(video and control) and provides a comprehensive interface for monitoring and '
         'controlling the robot.', 12, indent=1.27)

    para(doc, 'The dashboard layout consists of:', 12, indent=1.27)
    bullets(doc, [
        'Status Bar (top): Real-time indicators for camera connection, motor status, AI mode, '
        'and current date/time.',
        'Video Section (main): Live video feed with automotive-style overlays including '
        'perspective grid, corner brackets, crosshair with distance markers (0.5m, 1.0m, 2.0m), '
        'HUD panel (FPS, latency, mode, detections), navigation recommendation arrow, compass '
        'indicator, mini-map with obstacles, and quick stats.',
        'Control Panel (right): Mode toggle buttons (MANUAL/SEMI-AUTO), D-pad directional '
        'controls (Forward/Back/Left/Right/Stop), speed slider (0\u2013255 PWM), decision log '
        'showing last 10 AI decisions, and telemetry display (FPS, latency, obstacle count).',
    ])


# ---------------------------------------------------------------------------
# CHAPTER 4 - RESULTS AND TESTING
# ---------------------------------------------------------------------------

def add_chapter4(doc):
    chapter_break(doc, "4", "RESULTS AND TESTING")

    heading_section(doc, "4.1", "TESTING ENVIRONMENT SETUP")
    para(doc,
         'The DogBot Recon System was tested in both indoor laboratory environments and controlled '
         'outdoor settings. The testing environment and configuration used for the evaluation are '
         'described below:', 12, indent=1.27)
    bullets(doc, [
        'Hardware: AI-Thinker ESP32-CAM module, L298N motor driver with PWM, four N20 DC gear '
        'motors, 4WD chassis, 4\u00d7AA battery pack (6V).',
        'Backend Server: Laptop with Intel i5/i7 processor, 8GB+ RAM, running Windows 11, '
        'Python 3.10+, no dedicated GPU (CPU-only inference for YOLOv8).',
        'Network: WiFi 802.11n local network for ESP32 connection, HiveMQ Cloud MQTT broker '
        'for remote testing.',
        'Test Obstacles: Chairs, tables, boxes, bottles, books, wall corners, table legs, '
        'and human subjects at various distances.',
        'Dashboard Access: Chrome/Firefox browser at http://localhost:8000.',
        'Testing Modes: Both MANUAL (human-controlled) and SEMI_AUTO (AI-controlled) modes '
        'were evaluated.',
    ])

    img_box(doc, "4.1.1", "Assembled Hardware")

    para(doc,
         'The system was started using the recommended startup command (python run.py) to ensure '
         'proper Windows event loop configuration for MQTT compatibility. All services were '
         'verified to start correctly through the console log output and the /api/health endpoint.',
         12, indent=1.27)

    heading_section(doc, "4.2", "OBSTACLE DETECTION PERFORMANCE")
    para(doc,
         'The multi-sensor fusion CV pipeline was evaluated for detection accuracy at various '
         'distances and with different obstacle types. The six detection methods were tested both '
         'individually and in combination to assess the benefit of sensor fusion.', 12, indent=1.27)

    img_box(doc, "4.2.1", "Obstacle Detection Accuracy by Distance Range")

    para(doc,
         'The results demonstrate that multi-sensor fusion significantly improves detection rates '
         'compared to any single detection method, particularly in the critical DANGER zone '
         '(< 0.5m) where near-perfect detection is achieved. The slight increase in false positive '
         'rate is acceptable as the path planner treats unverified obstacles conservatively.',
         12, indent=1.27)

    para(doc,
         'The YOLOv8n-seg detector correctly identified common indoor objects including persons, '
         'tables, bottles, and backpacks with confidence scores above 0.5. The segmentation masks '
         'provided accurate object boundaries that improved grid accuracy compared to bounding '
         'boxes alone.', 12, indent=1.27)

    img_box(doc, "4.2.2", "Object Segmentation and Tracking")

    heading_section(doc, "4.3", "PATH PLANNING PERFORMANCE")
    para(doc,
         'The 5-stage path planner was evaluated for trajectory quality, planning latency, and '
         'navigation safety. The planner runs every frame and produces continuous steering and '
         'speed commands.', 12, indent=1.27)

    img_box(doc, "4.3.1", "Path Planner Performance Metrics")

    para(doc,
         'The path planner consistently achieved sub-2ms latency, making it suitable for '
         'real-time navigation at 15 FPS. The Kalman filter-based obstacle prediction proved '
         'valuable for avoiding moving obstacles, as the predictive cost map component shifted '
         'costs to future obstacle positions, causing the planner to preemptively steer away '
         'from approaching objects.', 12, indent=1.27)

    heading_section(doc, "4.4", "REAL-TIME FPS AND LATENCY MEASUREMENTS")
    para(doc,
         'The system\u2019s real-time performance was measured using the built-in telemetry system '
         'that reports FPS and per-frame latency through the dashboard and API.',
         12, indent=1.27)

    table_with_data(doc,
        headers=["Component", "Target Latency", "Measured Latency", "Execution Frequency"],
        rows=[
            ("ESP32 MJPEG Capture",  "~33 ms",   "30\u201340 ms",  "Every frame (30 FPS source)"),
            ("CV Pipeline Processing","< 30 ms",  "15\u201325 ms",  "Every frame"),
            ("ML Detection (YOLOv8)","< 60 ms",   "40\u201355 ms",  "Every 3 frames"),
            ("Path Planner",         "< 5 ms",    "~2 ms",          "Every frame"),
            ("AI Decision Engine",   "< 1 ms",    "< 1 ms",         "Every frame"),
            ("Frame Encoding (JPEG)","< 5 ms",    "3\u20135 ms",    "Every frame"),
            ("WebSocket Broadcast",  "< 2 ms",    "1\u20132 ms",    "Every frame"),
            ("Total Pipeline (local)","< 100 ms", "60\u201390 ms",  "15 FPS target"),
            ("MQTT Round-trip",      "< 500 ms",  "200\u2013400 ms","Per motor command"),
        ],
        col_widths=[2.0, 1.3, 1.5, 1.8],
        table_num="4.4.1",
        caption="System Performance Summary"
    )

    para(doc,
         'The system achieved a consistent 12\u201315 FPS processing rate on a laptop with an '
         'Intel i5 processor (no GPU). The frame-skipping strategy for ML detection (every '
         '3 frames) proved effective: running YOLOv8 on every frame would reduce the rate to '
         '8\u201310 FPS, while the skipping approach maintained high FPS with minimal detection '
         'delay.', 12, indent=1.27)

    heading_section(doc, "4.5", "MQTT REMOTE CONNECTIVITY TESTING")
    para(doc,
         'Remote connectivity was tested using HiveMQ Cloud as the MQTT broker with TLS '
         'encryption on port 8883. The following scenarios were evaluated:',
         12, indent=1.27)

    table_with_data(doc,
        headers=["Test Scenario", "Result", "Latency", "Notes"],
        rows=[
            ("Local WiFi MQTT",         "Pass", "50\u2013100 ms",   "ESP32 and backend on same network"),
            ("Remote MQTT (same city)",  "Pass", "150\u2013250 ms",  "Backend on different network"),
            ("Remote MQTT (different city)","Pass","200\u2013400 ms","Tested across locations"),
            ("TLS Encryption",          "Pass", "N/A",              "Verified encrypted connection"),
            ("Auto-reconnect",          "Pass", "~5 seconds",       "After simulated network drop"),
            ("Heartbeat monitoring",    "Pass", "15 sec timeout",   "Detected disconnection"),
            ("Motor command reliability","Pass","QoS 1",            "No lost commands observed"),
        ],
        col_widths=[2.0, 0.8, 1.3, 2.1],
        table_num="4.5.1",
        caption="MQTT Remote Connectivity Test Results"
    )

    para(doc,
         'The MQTT bridge demonstrated reliable remote connectivity with acceptable latency for '
         'manual control. The TLS encryption ensured secure communication, and the auto-reconnect '
         'mechanism handled network interruptions gracefully with a 5-second reconnection delay.',
         12, indent=1.27)

    heading_section(doc, "4.6", "DASHBOARD DEMONSTRATION")
    para(doc,
         'The web dashboard was tested for functionality, responsiveness, and real-time '
         'performance. All dashboard features were verified:', 12, indent=1.27)

    img_box(doc, "4.6.1", "Dashboard Demonstration")

    bullets(doc, [
        'Live video feed with obstacle detection overlays displayed smoothly at 12\u201315 FPS.',
        'D-pad controls sent motor commands with immediate response (<100 ms local latency).',
        'Speed slider adjusted motor PWM values correctly across the 0\u2013255 range.',
        'Mode switching between MANUAL and SEMI_AUTO worked seamlessly.',
    ])


# ---------------------------------------------------------------------------
# CHAPTER 5 - CONCLUSION AND FUTURE SCOPE
# ---------------------------------------------------------------------------

def add_chapter5(doc):
    chapter_break(doc, "5", "CONCLUSION AND FUTURE SCOPE")

    heading_section(doc, "5.1", "CONCLUSIONS")
    para(doc,
         'The DogBot Recon System has been successfully designed, developed, and tested as an '
         'autonomous reconnaissance robot with advanced AI-powered vision and path planning '
         'capabilities. The project demonstrates that sophisticated autonomous navigation can be '
         'achieved using affordable, commercially available hardware components combined with '
         'modern software engineering practices.', 12, indent=1.27)

    para(doc, 'The key achievements of this project are:', 12, indent=1.27)
    bullets(doc, [
        'Multi-Sensor Fusion Vision: A six-method obstacle detection pipeline (background '
        'subtraction, edge detection, gradient magnitude, optical flow, thin edge detection, '
        'and floor color analysis) achieves near-perfect detection rates in the critical danger '
        'zone (< 0.5m).',
        'Real-Time ML Detection: YOLOv8n-seg provides semantic object detection and instance '
        'segmentation at real-time frame rates on CPU hardware, with ByteTrack enabling '
        'persistent multi-object tracking across frames.',
        'Predictive Path Planning: A five-stage pipeline with Kalman filter-based obstacle state '
        'estimation achieves 2ms planning latency while evaluating 45 candidate trajectories '
        'using a five-component cost function.',
        'Intelligent Decision Making: A dual-layer AI architecture combines fast local planning '
        'with cloud LLM fallback, using multi-frame voting to ensure robust direction decisions.',
        'Worldwide Remote Access: MQTT over TLS enables secure motor control from any location, '
        'with automatic reconnection and heartbeat monitoring for reliability.',
        'Professional Web Dashboard: A feature-rich control interface provides live video with '
        'automotive-style overlays, real-time telemetry, manual and semi-automatic control modes.',
        'Modular Architecture: A service-oriented design with clear interfaces between components '
        'enables easy extension and modification of individual services.',
    ])

    heading_section(doc, "5.2", "FUTURE ENHANCEMENTS")
    bullets(doc, [
        'SLAM Integration: Implement visual SLAM (ORB-SLAM3 or OpenVSLAM) for persistent '
        'mapping and localization, enabling the robot to navigate to specific waypoints and '
        'build floor plans of explored areas.',
        'GPU Acceleration: Deploy on NVIDIA Jetson Nano or similar edge GPU platform to enable '
        'per-frame YOLOv8 detection and support larger models (YOLOv8s or YOLOv8m) for improved '
        'accuracy.',
        'MQTT Video Streaming: Implement compressed video streaming over MQTT using H.264 '
        'encoding for remote video access without VPN.',
        'Depth Sensor: Add Intel RealSense or similar depth sensor for accurate 3D obstacle '
        'mapping, replacing perspective-based distance estimation.',
        'Reinforcement Learning: Train a deep RL policy for navigation that can learn from '
        'experience and handle complex scenarios better than hand-tuned heuristics.',
        'Multi-Robot Coordination: Extend the MQTT architecture to support multiple DogBot '
        'units coordinating exploration of a shared environment.',
        'Voice Control: Integrate speech recognition for voice-commanded navigation.',
        'Custom YOLO Model: Train a specialized YOLOv8 model on domain-specific obstacles '
        '(specific indoor furniture, hazards) for improved detection accuracy.',
        'LiPo Battery with BMS: Upgrade to rechargeable LiPo batteries with battery management '
        'system for longer runtime and voltage monitoring.',
        'Autonomous Patrol Routes: Program the robot to autonomously patrol predefined routes, '
        'reporting anomalies detected along the way.',
        'Cloud Fleet Management: Build a cloud-based management platform for monitoring and '
        'controlling multiple DogBot units remotely.',
        'Robotic Arm Integration: Add a simple manipulator arm for physical interaction with '
        'the environment.',
    ])


# ---------------------------------------------------------------------------
# REFERENCES
# ---------------------------------------------------------------------------

def add_references(doc):
    doc.add_page_break()
    center(doc, "REFERENCES", 14, bold=True, sa=12)

    refs = [
        'J. Redmon, S. Divvala, R. Girshick, and A. Farhadi, \u201cYou Only Look Once: Unified, '
        'Real-Time Object Detection,\u201d in Proc. IEEE Conference on Computer Vision and Pattern '
        'Recognition (CVPR), 2016, pp. 779\u2013788.',

        'G. Jocher, A. Chaurasia, and J. Qiu, \u201cUltralytics YOLOv8,\u201d Ultralytics, 2023. '
        '[Online]. Available: https://github.com/ultralytics/ultralytics',

        'Y. Zhang, P. Sun, Y. Jiang, D. Yu, F. Weng, Z. Yuan, P. Luo, W. Liu, and X. Wang, '
        '\u201cByteTrack: Multi-Object Tracking by Associating Every Detection Box,\u201d in Proc. '
        'European Conference on Computer Vision (ECCV), 2022, pp. 1\u201321.',

        'D. Fox, W. Burgard, and S. Thrun, \u201cThe Dynamic Window Approach to Collision '
        'Avoidance,\u201d IEEE Robotics & Automation Magazine, vol. 4, no. 1, pp. 23\u201333, 1997.',

        'Z. Zivkovic, \u201cImproved Adaptive Gaussian Mixture Model for Background Subtraction,\u201d '
        'in Proc. International Conference on Pattern Recognition (ICPR), 2004, pp. 28\u201331.',

        'J. Canny, \u201cA Computational Approach to Edge Detection,\u201d IEEE Transactions on '
        'Pattern Analysis and Machine Intelligence, vol. PAMI-8, no. 6, pp. 679\u2013698, 1986.',

        'G. Farneback, \u201cTwo-Frame Motion Estimation Based on Polynomial Expansion,\u201d in '
        'Proc. Scandinavian Conference on Image Analysis (SCIA), 2003, pp. 363\u2013370.',

        'P. E. Hart, N. J. Nilsson, and B. Raphael, \u201cA Formal Basis for the Heuristic '
        'Determination of Minimum Cost Paths,\u201d IEEE Transactions on Systems Science and '
        'Cybernetics, vol. 4, no. 2, pp. 100\u2013107, 1968.',

        'R. A. Brooks, \u201cA Robust Layered Control System for a Mobile Robot,\u201d IEEE Journal '
        'on Robotics and Automation, vol. 2, no. 1, pp. 14\u201323, 1986.',

        'R. E. Kalman, \u201cA New Approach to Linear Filtering and Prediction Problems,\u201d '
        'Journal of Basic Engineering, vol. 82, no. 1, pp. 35\u201345, 1960.',

        'Espressif Systems, \u201cESP32 Technical Reference Manual,\u201d 2023. [Online]. Available: '
        'https://www.espressif.com/en/products/socs/esp32',

        'A. Banks and R. Gupta, \u201cMQTT Version 5.0 - OASIS Standard,\u201d OASIS, 2019. '
        '[Online]. Available: https://docs.oasis-open.org/mqtt/mqtt/v5.0/',

        'S. Ramirez, \u201cFastAPI - Modern, Fast Web Framework for Building APIs with Python,\u201d '
        '2019. [Online]. Available: https://fastapi.tiangolo.com/',

        'G. Bradski, \u201cThe OpenCV Library,\u201d Dr. Dobb\u2019s Journal of Software Tools, 2000.',

        'HiveMQ, \u201cHiveMQ Cloud - Fully Managed MQTT Platform,\u201d 2023. [Online]. Available: '
        'https://www.hivemq.com/mqtt-cloud-broker/',

        'S. M. LaValle, \u201cRapidly-Exploring Random Trees: A New Tool for Path Planning,\u201d '
        'Technical Report TR 98-11, Computer Science Dept., Iowa State University, 1998.',

        'J. Borenstein and Y. Koren, \u201cThe Vector Field Histogram - Fast Obstacle Avoidance '
        'for Mobile Robots,\u201d IEEE Transactions on Robotics and Automation, vol. 7, no. 3, '
        'pp. 278\u2013288, 1991.',

        'S. Thrun, W. Burgard, and D. Fox, \u201cProbabilistic Robotics,\u201d MIT Press, 2005.',

        'AI-Thinker, \u201cESP32-CAM Development Board Datasheet,\u201d 2020. [Online]. Available: '
        'https://docs.ai-thinker.com/en/esp32-cam',

        'STMicroelectronics, \u201cL298N Dual Full-Bridge Driver Datasheet,\u201d 2000. [Online]. '
        'Available: https://www.st.com/resource/en/datasheet/l298.pdf',
    ]

    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(6)
        p.paragraph_format.line_spacing = Pt(20)
        p.paragraph_format.left_indent  = Cm(1.0)
        p.paragraph_format.first_line_indent = Cm(-1.0)
        r_num = p.add_run(f"{i}. ")
        r_num.bold = True; r_num.font.size = Pt(11); r_num.font.name = TNR
        r_txt = p.add_run(ref)
        r_txt.font.size = Pt(11); r_txt.font.name = TNR


# ---------------------------------------------------------------------------
# APPENDIX - PO & PSO ATTAINMENT
# ---------------------------------------------------------------------------

def add_po_pso(doc):
    doc.add_page_break()
    center(doc, "APPENDIX", 14, bold=True, sa=6)
    center(doc, "LEARNING OUTCOMES", 12, bold=True, sa=12)
    center(doc, "PROGRAM OUTCOME ATTAINMENT", 13, bold=True, sa=10)

    po_rows = [
        ("PO 1",  "Engineering knowledge",
         "Yes",
         "Applied robotics, electronics, computer vision, and wireless communication principles "
         "in developing a comprehensive surveillance system integrating multiple engineering domains."),
        ("PO 2",  "Problem analysis",
         "Yes",
         "Analyzed security challenges in surveillance operations and developed a robotic solution "
         "addressing personnel safety while maintaining operational effectiveness."),
        ("PO 3",  "Design/Development of solutions",
         "Yes",
         "Designed and implemented a complete robotic surveillance platform incorporating quadruped "
         "locomotion, wireless control, and AI-powered detection capabilities."),
        ("PO 4",  "Conduct investigations of complex problems",
         "Yes",
         "Systematically investigated machine learning model performance, wireless communication "
         "reliability, and robot navigation effectiveness in challenging environments."),
        ("PO 5",  "Modern Tool usage",
         "Yes",
         "Utilized Arduino IDE, Python ML frameworks, mobile app development tools, and computer "
         "vision libraries for comprehensive system development."),
        ("PO 6",  "The Engineer and society",
         "Yes",
         "Developed technology enhancing safety capabilities while prioritizing responsible and "
         "public security in operational scenarios."),
        ("PO 7",  "Environment and Sustainability",
         "Yes",
         "Created an energy-efficient surveillance system with minimal environmental impact and "
         "sustainable operation."),
        ("PO 8",  "Ethics",
         "Yes",
         "Ensured responsible development of surveillance technology with appropriate privacy "
         "considerations and ethical deployment guidelines."),
        ("PO 9",  "Individual and team work",
         "Yes",
         "Successfully collaborated across mechanical design, electronics integration, software "
         "development, and testing phases with effective task distribution and coordination."),
        ("PO 10", "Communication",
         "Yes",
         "Documented complex technical concepts clearly and interfaced effectively with potential "
         "end users to gather requirements and validate system functionality."),
        ("PO 11", "Project management and finance",
         "Yes",
         "Completed project within allocated timeframe and budget constraints through strategic "
         "component selection and efficient development methodologies."),
        ("PO 12", "Life-long learning",
         "Yes",
         "Acquired new expertise in quadruped robotics, computer vision, machine learning "
         "deployment, and specialized surveillance applications throughout project development."),
    ]

    table_with_data(doc,
        headers=["PO. No", "Graduate Attribute", "Attained", "Justification"],
        rows=po_rows,
        col_widths=[0.7, 1.5, 0.8, 3.5],
    )

    doc.add_paragraph()
    center(doc, "PROGRAM SPECIFIC OUTCOMES ATTAINMENT", 13, bold=True, sa=10)

    pso_rows = [
        ("PSO 1",
         "To analyse, design and develop solutions by applying the concepts of Robotics for "
         "societal and industrial needs.",
         "Yes",
         "Created a robotic surveillance system that addresses critical societal security needs "
         "while demonstrating practical applications of advanced robotics concepts in law "
         "enforcement and public safety."),
        ("PSO 2",
         "To create innovative ideas and solutions for real time problems in Manufacturing sector "
         "by adapting the automation tools and technologies.",
         "Yes",
         "Developed an innovative surveillance automation solution that can be adapted for "
         "security manufacturing environments, perimeter monitoring, and automated inspection "
         "systems using real-time AI processing and robotic mobility."),
    ]

    table_with_data(doc,
        headers=["PSO. No", "Graduate Attribute", "Attained", "Justification"],
        rows=pso_rows,
        col_widths=[0.7, 2.0, 0.8, 3.0],
    )


# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------

def main():
    doc = Document()

    # Page setup: A4, 1-inch margins
    section = doc.sections[0]
    section.page_width  = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.0)

    # Default style
    style = doc.styles['Normal']
    style.font.name = TNR
    style.font.size = Pt(12)

    print("Building report sections ...")

    add_title_page(doc)
    print("  [1/11] Title page")

    add_bonafide_certificate(doc)
    print("  [2/11] Bonafide certificate")

    add_acknowledgement(doc)
    print("  [3/11] Acknowledgement")

    add_abstract(doc)
    print("  [4/11] Abstract")

    add_table_of_contents(doc)
    print("  [5/11] Table of contents")

    add_list_of_figures(doc)
    print("  [6/11] List of figures")

    add_list_of_tables(doc)
    print("  [7/11] List of tables")

    add_chapter1(doc)
    print("  [8/11] Chapter 1 \u2013 Introduction")

    add_chapter2(doc)
    print("  [9/11] Chapter 2 \u2013 Literature Survey")

    add_chapter3(doc)
    print("  [10/11] Chapter 3 \u2013 Methodology (incl. Hardware & Software)")

    add_chapter4(doc)
    print("  [11/11] Chapter 4 \u2013 Results & Testing")

    add_chapter5(doc)
    print("  [12/12] Chapter 5 \u2013 Conclusion & Future Scope")

    add_references(doc)
    print("  [13/13] References")

    add_po_pso(doc)
    print("  [14/14] PO & PSO Attainment")

    out_path = os.path.join(os.path.dirname(__file__), "22MT060_22MT022_REPORT.docx")
    doc.save(out_path)
    print(f"\nReport saved to: {out_path}")


if __name__ == "__main__":
    main()
